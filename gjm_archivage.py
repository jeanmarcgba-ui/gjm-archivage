#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
GJM TECNOLOGIE - Application d'Archivage des Dossiers
Interface moderne avec CustomTkinter
"""

# ── Version de l'application ──────────────────────────────────────
VERSION         = "1.0.0"
GITHUB_USER     = "jeanmarcgba-ui"          # Votre nom d'utilisateur GitHub
GITHUB_REPO     = "gjm-archivage"           # Nom du dépôt GitHub
GITHUB_BRANCH   = "main"                    # Branche principale
# URL du fichier version sur GitHub (fichier texte contenant juste le numéro)
URL_VERSION     = f"https://raw.githubusercontent.com/{GITHUB_USER}/{GITHUB_REPO}/{GITHUB_BRANCH}/version.txt"
# URL du fichier application sur GitHub
URL_APP         = f"https://raw.githubusercontent.com/{GITHUB_USER}/{GITHUB_REPO}/{GITHUB_BRANCH}/gjm_archivage.py"

import os, sys, json, shutil, subprocess, platform, datetime, hashlib, sqlite3
from pathlib import Path

try:
    import customtkinter as ctk
    from customtkinter import CTkFont
    CTK = True
except ImportError:
    CTK = False
    import tkinter as tk
    from tkinter import ttk

import tkinter as tk
from tkinter import ttk, messagebox, filedialog, simpledialog

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    EXCEL_OK = True
except: EXCEL_OK = False

try:
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.enums import TA_CENTER
    PDF_OK = True
except: PDF_OK = False

try:
    from docx import Document as DocxDoc
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except: DOCX_OK = False

# ── Constantes ────────────────────────────────────────────────────────────────
APP_DIR  = Path(os.path.dirname(os.path.abspath(__file__)))
DB_PATH  = APP_DIR / "gjm_data.db"
DOCS_DIR = APP_DIR / "dossiers_clients"

# ── Connexion réseau robuste (WAL + retry) ────────────────────────────────────
import time, threading
_db_lock = threading.Lock()

def get_conn(timeout=10):
    """Connexion SQLite optimisée multi-utilisateurs réseau."""
    conn = sqlite3.connect(str(DB_PATH), timeout=timeout, check_same_thread=False)
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA synchronous=NORMAL")
    conn.execute("PRAGMA busy_timeout=10000")
    conn.execute("PRAGMA cache_size=-4000")
    conn.row_factory = sqlite3.Row
    return conn

def db_execute(sql, params=(), fetchone=False, fetchall=False, lastrowid=False):
    """Exécute une requête SQL avec retry automatique (réseau partagé)."""
    for attempt in range(5):
        try:
            with _db_lock:
                conn = get_conn()
                c = conn.cursor()
                c.execute(sql, params)
                conn.commit()
                if fetchone:  r = c.fetchone();  conn.close(); return dict(r) if r else None
                if fetchall:  r = c.fetchall();  conn.close(); return [dict(x) for x in r]
                if lastrowid: r = c.lastrowid;   conn.close(); return r
                conn.close(); return True
        except sqlite3.OperationalError as e:
            if "locked" in str(e).lower() and attempt < 4:
                time.sleep(0.3 * (attempt + 1))
                continue
            raise
    return None

# Palette moderne
C_BG       = "#1a1a2e"   # Fond principal sombre
C_BG2      = "#16213e"   # Fond secondaire
C_PANEL    = "#0f3460"   # Panneaux
C_ACCENT   = "#e94560"   # Accent rouge/rose
C_BLUE     = "#533483"   # Bleu violet
C_GREEN    = "#00b4d8"   # Bleu cyan
C_WHITE    = "#e0e0e0"   # Texte principal
C_GOLD     = "#ffd700"   # Or
C_CARD     = "#1e2a4a"   # Carte
C_SUCCESS  = "#06d6a0"   # Succès vert
C_WARNING  = "#ffd166"   # Avertissement
C_HOVER    = "#2a3a6a"   # Survol

# ── Utilisateurs dynamiques (fichier JSON local) ─────────────────
USERS_FILE = APP_DIR / "gjm_users.json"

def charger_users():
    if USERS_FILE.exists():
        try:
            data = json.loads(USERS_FILE.read_text(encoding="utf-8"))
            if data: return data
        except: pass
    return {}

def sauvegarder_users(d):
    USERS_FILE.write_text(json.dumps(d, ensure_ascii=False, indent=2), encoding="utf-8")

USERS = charger_users()

def hash_pwd(p): return hashlib.sha256(p.encode()).hexdigest()

# ── Compte super-administrateur GJM Tecnologie (créateur) ───────────────
# Ce compte permet à GJM Tecnologie d'accéder à toute installation client
# en cas de problème technique. Invisible pour les clients.
_SA_U = "635a4747c4183436e5fa955fb9a45214126487c414e0d7e48131957a0e7658d5"  # hash de l'identifiant
_SA_P = "edd7c672fae2698d30e8f4bb9ab027b9f3f2d1e0daf167413d1608fabf241140"  # hash du mot de passe

def _est_super_admin(nom, mdp):
    """Vérifie si c'est le compte super-admin GJM Tecnologie."""
    return (hash_pwd(nom) == _SA_U and hash_pwd(mdp) == _SA_P)


# ── Générateur de clé de licence ────────────────────────────────
def generer_cle_licence(duree_jours, client_id=""):
    """Génère une clé de licence encodée (pour le vendeur GJM Tecnologie)."""
    import base64
    fin = (datetime.date.today() + datetime.timedelta(days=duree_jours)).strftime("%Y-%m-%d")
    payload = f"{fin}|{client_id}|GJM2025"
    encoded = base64.b64encode(payload.encode()).decode()
    # Clé lisible par blocs de 4
    cle = encoded.replace("=","").upper()
    return "-".join([cle[i:i+4] for i in range(0, min(len(cle),16), 4)])

def valider_cle_licence(cle):
    """Valide une clé de licence et retourne la date de fin ou None."""
    import base64
    try:
        cle_clean = cle.replace("-","").replace(" ","").upper()
        # Padding base64
        padding = 4 - len(cle_clean) % 4
        if padding != 4: cle_clean += "=" * padding
        decoded = base64.b64decode(cle_clean).decode()
        parts = decoded.split("|")
        if len(parts) >= 3 and parts[2] == "GJM2025":
            return parts[0]  # date fin YYYY-MM-DD
        return None
    except:
        return None

def statut_licence():
    """Retourne (texte, couleur) selon l etat de la licence."""
    fin = get_param("duree_fin", "")
    if not fin:
        return ("⚠  Licence non definie", C_WARNING)
    try:
        fin_date = datetime.datetime.strptime(fin, "%Y-%m-%d").date()
        today    = datetime.date.today()
        jours    = (fin_date - today).days
        fin_fmt  = fin_date.strftime("%d/%m/%Y")
        if jours < 0:
            return ("⛔  Expiree depuis " + str(abs(jours)) + "j – " + fin_fmt, C_ACCENT)
        elif jours == 0:
            return ("⚠  Expire aujourd hui – " + fin_fmt, C_WARNING)
        elif jours <= 15:
            return ("⚠  Expire dans " + str(jours) + "j – " + fin_fmt, C_WARNING)
        else:
            return ("✅  Valide – " + fin_fmt + "  (" + str(jours) + " jours)", C_SUCCESS)
    except:
        return ("⚠  Date invalide", C_WARNING)

def ouvrir_fichier(chemin):
    try:
        if platform.system() == "Windows": os.startfile(chemin)
        elif platform.system() == "Darwin": subprocess.call(["open", chemin])
        else: subprocess.call(["xdg-open", chemin])
    except Exception as e:
        messagebox.showerror("Erreur", f"Impossible d'ouvrir :\n{e}")

def init_db():
    conn = get_conn()
    c = conn.cursor()
    c.executescript("""
    CREATE TABLE IF NOT EXISTS parametres (cle TEXT PRIMARY KEY, valeur TEXT);
    CREATE TABLE IF NOT EXISTS dossiers_entrants (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        numero_dossier TEXT UNIQUE, nom_prenom TEXT, adresse TEXT,
        telephone TEXT, email TEXT, ilot_numero TEXT, lot_numero TEXT,
        superficie TEXT, nom_lotissement TEXT, commune TEXT,
        nature_dossier TEXT, imputer_a TEXT, date_entree TEXT,
        personne_contact TEXT, telephone_contact TEXT,
        mois INTEGER, annee INTEGER,
        date_creation TEXT DEFAULT (datetime('now','localtime'))
    );
    CREATE TABLE IF NOT EXISTS documents_clients (
        id INTEGER PRIMARY KEY AUTOINCREMENT, dossier_id INTEGER,
        nom_fichier TEXT, chemin_fichier TEXT, type_fichier TEXT,
        date_ajout TEXT DEFAULT (datetime('now','localtime')),
        FOREIGN KEY (dossier_id) REFERENCES dossiers_entrants(id)
    );
    CREATE TABLE IF NOT EXISTS dossiers_sortants (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        numero_dossier TEXT, nom_prenom TEXT, telephone TEXT,
        adresse TEXT, date_sortie TEXT, mois INTEGER, annee INTEGER,
        date_creation TEXT DEFAULT (datetime('now','localtime'))
    );
    """)
    conn.commit(); conn.close()

def get_param(k, d=""): 
    conn=get_conn(); c=conn.cursor()
    c.execute("SELECT valeur FROM parametres WHERE cle=?",(k,)); r=c.fetchone(); conn.close()
    return r[0] if r else d

def set_param(k,v): 
    conn=get_conn(); c=conn.cursor()
    c.execute("INSERT OR REPLACE INTO parametres VALUES(?,?)",(k,v)); conn.commit(); conn.close()

def prochain_numero():
    conn=get_conn(); c=conn.cursor()
    c.execute("SELECT COUNT(*) FROM dossiers_entrants"); n=c.fetchone()[0]+1; conn.close()
    return f"DE{n:03d}"

# ══════════════════════════════════════════════════════════════════════════════
# COMPOSANTS UI MODERNES
# ══════════════════════════════════════════════════════════════════════════════

def make_card(parent, **kw):
    f = tk.Frame(parent, bg=C_CARD, bd=0, highlightthickness=1,
                 highlightbackground=C_PANEL, **kw)
    return f

def make_btn(parent, text, cmd, color=C_ACCENT, fg=C_WHITE, w=None, **kw):
    b = tk.Button(parent, text=text, command=cmd, bg=color, fg=fg,
                  font=("Segoe UI", 11, "bold"), relief="flat", cursor="hand2",
                  activebackground=color, activeforeground=fg,
                  padx=18, pady=9, bd=0, **kw)
    if w: b.config(width=w)
    b.bind("<Enter>", lambda e: b.config(bg=_lighten(color)))
    b.bind("<Leave>", lambda e: b.config(bg=color))
    return b

def _lighten(hex_color):
    """Éclaircit légèrement une couleur hex."""
    try:
        r = min(255, int(hex_color[1:3],16)+30)
        g = min(255, int(hex_color[3:5],16)+30)
        b = min(255, int(hex_color[5:7],16)+30)
        return f"#{r:02x}{g:02x}{b:02x}"
    except: return hex_color

def make_label(parent, text, size=11, bold=False, color=C_WHITE, **kw):
    return tk.Label(parent, text=text, bg=C_CARD if kw.get('card') else C_BG,
                    fg=color, font=("Segoe UI", size, "bold" if bold else "normal"), **kw)

def make_entry(parent, var, width=28, show=None):
    e = tk.Entry(parent, textvariable=var, font=("Segoe UI", 12),
                 bg="#1e2a4a", fg=C_WHITE, insertbackground=C_WHITE,
                 relief="flat", bd=0, highlightthickness=1,
                 highlightbackground=C_PANEL, highlightcolor=C_GREEN, width=width)
    if show: e.config(show=show)
    return e

def style_treeview():
    style = ttk.Style()
    style.theme_use("clam")
    style.configure("Modern.Treeview",
        background=C_CARD, foreground=C_WHITE,
        fieldbackground=C_CARD, rowheight=34,
        font=("Segoe UI", 11), borderwidth=0)
    style.configure("Modern.Treeview.Heading",
        background=C_PANEL, foreground=C_GOLD,
        font=("Segoe UI", 11, "bold"), relief="flat", borderwidth=0)
    style.map("Modern.Treeview",
        background=[("selected", C_ACCENT)],
        foreground=[("selected", C_WHITE)])
    style.map("Modern.Treeview.Heading",
        background=[("active", C_BLUE)])

    style.configure("Tab.TNotebook", background=C_BG2, borderwidth=0, tabmargins=0)
    style.configure("Tab.TNotebook.Tab",
        background=C_BG2, foreground="#aaaaaa",
        font=("Segoe UI", 11, "bold"), padding=[18, 9])
    style.map("Tab.TNotebook.Tab",
        background=[("selected", C_PANEL)],
        foreground=[("selected", C_GOLD)])

    style.configure("TScrollbar", background=C_PANEL, troughcolor=C_BG2,
                    arrowcolor=C_WHITE, borderwidth=0)
    style.configure("TCombobox",
        fieldbackground=C_CARD, background=C_PANEL,
        foreground=C_WHITE, selectbackground=C_ACCENT,
        font=("Segoe UI", 12))

# ══════════════════════════════════════════════════════════════════════════════
# FENETRE CONNEXION
# ══════════════════════════════════════════════════════════════════════════════
# ═════════════════════════════════════════════════════════════════════════════
# FENETRE MISE A JOUR
# ═════════════════════════════════════════════════════════════════════════════
class FenetreMiseAJour(tk.Toplevel):
    def __init__(self, parent, nouvelle_version):
        super().__init__(parent)
        self.parent           = parent
        self.nouvelle_version = nouvelle_version
        self.title("Mise a jour disponible – GJM Archivage")
        self.configure(bg=C_BG)
        self.resizable(False, False)
        self.grab_set()
        w, h = 500, 320
        x = (self.winfo_screenwidth()-w)//2
        y = (self.winfo_screenheight()-h)//2
        self.geometry(f"{w}x{h}+{x}+{y}")
        self._build()

    def _build(self):
        hdr = tk.Frame(self, bg=C_SUCCESS, pady=15)
        hdr.pack(fill="x")
        tk.Label(hdr, text="Nouvelle version disponible !",
                 font=("Segoe UI Black", 14, "bold"),
                 bg=C_SUCCESS, fg="#1a1a2e").pack()

        body = tk.Frame(self, bg=C_BG, padx=35, pady=20)
        body.pack(fill="both", expand=True)

        cadre_v = tk.Frame(body, bg=C_CARD, padx=20, pady=15,
                           highlightthickness=1, highlightbackground=C_PANEL)
        cadre_v.pack(fill="x", pady=(0,15))

        tk.Label(cadre_v, text="Version actuelle : " + VERSION,
                 font=("Segoe UI", 12), bg=C_CARD, fg=C_ACCENT, anchor="w").pack(fill="x")
        tk.Label(cadre_v, text="Nouvelle version : " + self.nouvelle_version,
                 font=("Segoe UI", 12, "bold"), bg=C_CARD, fg=C_SUCCESS, anchor="w").pack(fill="x", pady=(8,0))

        tk.Label(body, text="La mise a jour prend moins de 30 secondes.",
                 font=("Segoe UI", 11), bg=C_BG, fg=C_WHITE).pack(pady=(0,4))
        tk.Label(body, text="Vos donnees et votre licence sont conservees.",
                 font=("Segoe UI", 11), bg=C_BG, fg=C_WHITE).pack(pady=(0,15))

        self.lbl_prog = tk.Label(body, text="",
                                  font=("Segoe UI", 11, "bold"), bg=C_BG, fg=C_GOLD)
        self.lbl_prog.pack(pady=5)

        btn_f = tk.Frame(body, bg=C_BG); btn_f.pack(fill="x")
        make_btn(btn_f, "Mettre a jour maintenant",
                 self._installer, color=C_SUCCESS, fg="#1a1a2e").pack(
                 side="left", fill="x", expand=True, padx=(0,5), ipady=6)
        make_btn(btn_f, "Plus tard", self.destroy,
                 color="#333355", fg="#aaaaaa").pack(side="left", ipady=6, padx=(5,0))

    def _installer(self):
        self.lbl_prog.config(text="Telechargement en cours...")
        self.update()
        succes = telecharger_mise_a_jour(self.lbl_prog)
        if succes:
            self.lbl_prog.config(text="Mise a jour reussie !", fg=C_SUCCESS)
            self.update()
            msg = ("Application mise a jour vers la version " +
                   self.nouvelle_version + "\n\nL'application va redemarrer.")
            messagebox.showinfo("Mise a jour reussie", msg, parent=self)
            self.destroy()
            python = sys.executable
            os.execl(python, python, *sys.argv)
        else:
            self.lbl_prog.config(text="Echec. Verifiez votre connexion.", fg=C_ACCENT)


# ═════════════════════════════════════════════════════════════════════════════
# FENETRE RENOUVELLEMENT DE LICENCE
# ═════════════════════════════════════════════════════════════════════════════
class FenetreRenouvellement(tk.Toplevel):
    """Fenêtre de saisie de la nouvelle clé de licence."""

    def __init__(self, parent, callback=None, expiree=False):
        super().__init__(parent)
        self.parent   = parent
        self.callback = callback
        self.expiree  = expiree
        self.title("Renouvellement de licence – GJM Archivage")
        self.configure(bg=C_BG)
        self.resizable(False, False)
        self.grab_set()
        # Si expirée, empêcher de fermer sans saisir une clé
        if expiree:
            self.protocol("WM_DELETE_WINDOW", lambda: None)
        w, h = 520, 420
        x = (self.winfo_screenwidth()-w)//2
        y = (self.winfo_screenheight()-h)//2
        self.geometry(f"{w}x{h}+{x}+{y}")
        self._build()

    def _build(self):
        # En-tête coloré selon le contexte
        hdr_col = C_ACCENT if self.expiree else C_PANEL
        hdr = tk.Frame(self, bg=hdr_col, pady=18)
        hdr.pack(fill="x")

        if self.expiree:
            tk.Label(hdr, text="⛔  Licence Expirée",
                     font=("Segoe UI Black", 16, "bold"),
                     bg=hdr_col, fg=C_WHITE).pack()
            tk.Label(hdr, text="Entrez une nouvelle clé pour continuer à utiliser l'application",
                     font=("Segoe UI", 10), bg=hdr_col, fg="#ffcccc").pack(pady=(4,0))
        else:
            tk.Label(hdr, text="🔑  Renouveler la Licence",
                     font=("Segoe UI Black", 16, "bold"),
                     bg=hdr_col, fg=C_GOLD).pack()
            txt, col = statut_licence()
            tk.Label(hdr, text="Statut actuel : " + txt,
                     font=("Segoe UI", 10), bg=hdr_col, fg=col).pack(pady=(4,0))

        body = tk.Frame(self, bg=C_BG, padx=35, pady=25)
        body.pack(fill="both", expand=True)

        # Clé actuelle
        fin = get_param("duree_fin", "")
        if fin:
            try:
                fin_fmt = datetime.datetime.strptime(fin, "%Y-%m-%d").strftime("%d/%m/%Y")
                jours   = (datetime.datetime.strptime(fin, "%Y-%m-%d").date() - datetime.date.today()).days
                if jours < 0:
                    info = "Expirée depuis " + str(abs(jours)) + " jour(s) – " + fin_fmt
                    info_col = C_ACCENT
                else:
                    info = "Valide jusqu'au " + fin_fmt + " (" + str(jours) + " jours restants)"
                    info_col = C_SUCCESS
                tk.Label(body, text=info, font=("Segoe UI", 11, "bold"),
                         bg=C_BG, fg=info_col).pack(anchor="w", pady=(0,15))
            except: pass

        tk.Label(body, text="Nouvelle clé de licence :",
                 font=("Segoe UI", 12, "bold"), bg=C_BG, fg=C_GREEN, anchor="w").pack(fill="x", pady=(0,5))
        tk.Label(body, text="Fournie par GJM Tecnologie – Format : XXXX-XXXX-XXXX-XXXX",
                 font=("Segoe UI", 10), bg=C_BG, fg="#aaaaaa", anchor="w").pack(fill="x")

        self.var_cle = tk.StringVar()
        make_entry(body, self.var_cle, width=38).pack(fill="x", ipady=8, pady=(6,4))

        self.lbl_statut = tk.Label(body, text="",
                                    font=("Segoe UI", 11, "bold"), bg=C_BG, fg=C_WARNING)
        self.lbl_statut.pack(anchor="w", pady=5)

        # Boutons
        make_btn(body, "Valider la nouvelle cle",
                 self._valider, color=C_SUCCESS, fg="#1a1a2e").pack(fill="x", pady=6, ipady=8)

        if not self.expiree:
            make_btn(body, "Annuler", self.destroy,
                     color="#333355", fg="#aaaaaa").pack(fill="x", ipady=5)

        tk.Label(body,
                 text="Contactez GJM Tecnologie pour obtenir une nouvelle cle de licence.",
                 font=("Segoe UI", 9, "italic"), bg=C_BG, fg="#555").pack(pady=(15,0))

    def _valider(self):
        cle = self.var_cle.get().strip()
        if not cle:
            self.lbl_statut.config(text="Entrez une cle de licence.", fg=C_WARNING)
            return

        fin_date = valider_cle_licence(cle)
        if not fin_date:
            self.lbl_statut.config(
                text="Cle invalide. Verifiez et reessayez.", fg=C_ACCENT)
            return

        # Vérifier que la nouvelle licence est dans le futur
        try:
            fin = datetime.datetime.strptime(fin_date, "%Y-%m-%d").date()
            jours = (fin - datetime.date.today()).days
            if jours < 0:
                self.lbl_statut.config(
                    text="Cette cle est deja expiree (" + str(abs(jours)) + " jours). Demandez une nouvelle cle.", fg=C_ACCENT)
                return
            fin_fmt = fin.strftime("%d/%m/%Y")
        except:
            self.lbl_statut.config(text="Erreur dans la cle.", fg=C_ACCENT)
            return

        # Sauvegarder la nouvelle date
        set_param("duree_fin", fin_date)
        self.lbl_statut.config(
            text="Licence renouvelee ! Valide jusqu'au " + fin_fmt + " (" + str(jours) + " jours)",
            fg=C_SUCCESS)

        # Mettre à jour le badge topbar si possible
        if self.callback:
            try: self.callback()
            except: pass

        self.after(1500, self.destroy)


# ═════════════════════════════════════════════════════════════════════════════
# ASSISTANT PREMIER DEMARRAGE
# ═════════════════════════════════════════════════════════════════════════════
class AssistantPremierDemarrage(tk.Toplevel):
    """Affiché au tout premier lancement — configure l'application pour le client."""

    def __init__(self, parent, callback):
        super().__init__(parent)
        self.parent   = parent
        self.callback = callback
        self.title("Configuration initiale – GJM Archivage")
        self.configure(bg=C_BG)
        self.resizable(False, False)
        self.grab_set()
        self.protocol("WM_DELETE_WINDOW", lambda: None)  # Interdit de fermer
        w, h = 560, 660
        x = (self.winfo_screenwidth()-w)//2
        y = (self.winfo_screenheight()-h)//2
        self.geometry(f"{w}x{h}+{x}+{y}")
        self._build()

    def _build(self):
        # En-tête
        hdr = tk.Frame(self, bg=C_PANEL, pady=18)
        hdr.pack(fill="x")
        tk.Label(hdr, text="GJM TECNOLOGIE",
                 font=("Segoe UI Black", 18, "bold"),
                 bg=C_PANEL, fg=C_GOLD).pack()
        tk.Label(hdr, text="Configuration initiale de votre application",
                 font=("Segoe UI", 11), bg=C_PANEL, fg="#aaaaaa").pack(pady=(4,0))

        # Corps scrollable
        canvas = tk.Canvas(self, bg=C_BG, highlightthickness=0)
        sb = ttk.Scrollbar(self, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=sb.set)
        sb.pack(side="right", fill="y")
        canvas.pack(fill="both", expand=True)
        body = tk.Frame(canvas, bg=C_BG, padx=35, pady=20)
        win  = canvas.create_window((0,0), window=body, anchor="nw")
        canvas.bind("<Configure>", lambda e: canvas.itemconfig(win, width=e.width))
        body.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

        tk.Label(body, text="Bienvenue ! Configurez votre application en 3 etapes.",
                 font=("Segoe UI", 11), bg=C_BG, fg=C_WHITE).pack(anchor="w", pady=(0,15))

        # ── ETAPE 1 : Clé de licence ─────────────────────────────────
        tk.Frame(body, bg=C_GOLD, height=2).pack(fill="x", pady=(0,8))
        tk.Label(body, text="1  Cle de licence",
                 font=("Segoe UI Black", 12), bg=C_BG, fg=C_GOLD).pack(anchor="w")
        tk.Label(body, text="Entrez la cle fournie par GJM Tecnologie :",
                 font=("Segoe UI", 10), bg=C_BG, fg="#aaaaaa").pack(anchor="w", pady=(2,4))
        self.var_cle = tk.StringVar()
        e_cle = make_entry(body, self.var_cle, width=38)
        e_cle.pack(fill="x", ipady=6)
        tk.Label(body, text="Format : XXXX-XXXX-XXXX-XXXX",
                 font=("Segoe UI", 9, "italic"), bg=C_BG, fg="#666").pack(anchor="w")
        self.lbl_cle_statut = tk.Label(body, text="", font=("Segoe UI", 10, "bold"),
                                        bg=C_BG, fg=C_WARNING)
        self.lbl_cle_statut.pack(anchor="w", pady=2)

        # ── ETAPE 2 : Nom entreprise ─────────────────────────────────
        tk.Frame(body, bg=C_BLUE, height=2).pack(fill="x", pady=(12,8))
        tk.Label(body, text="2  Nom de votre entreprise",
                 font=("Segoe UI Black", 12), bg=C_BG, fg=C_GREEN).pack(anchor="w")
        self.var_nom_ent = tk.StringVar()
        make_entry(body, self.var_nom_ent, width=38).pack(fill="x", ipady=6, pady=(4,0))

        # ── ETAPE 3 : Compte administrateur ──────────────────────────
        tk.Frame(body, bg=C_ACCENT, height=2).pack(fill="x", pady=(12,8))
        tk.Label(body, text="3  Compte administrateur",
                 font=("Segoe UI Black", 12), bg=C_BG, fg=C_ACCENT).pack(anchor="w")
        tk.Label(body, text="Cet identifiant vous permettra de gerer l'application.",
                 font=("Segoe UI", 10), bg=C_BG, fg="#aaaaaa").pack(anchor="w", pady=(2,6))

        grille = tk.Frame(body, bg=C_BG); grille.pack(fill="x")
        for i, (lbl, key, show) in enumerate([
            ("Nom d'utilisateur", "admin_nom", None),
            ("Mot de passe",      "admin_pwd", "●"),
            ("Confirmer MDP",     "admin_pwd2","●"),
        ]):
            tk.Label(grille, text=lbl+":", font=("Segoe UI",10,"bold"),
                     bg=C_BG, fg=C_WHITE, width=18, anchor="e").grid(row=i,column=0,sticky="e",pady=5)
            var = tk.StringVar()
            setattr(self, "var_"+key, var)
            make_entry(grille, var, width=24, show=show).grid(row=i,column=1,padx=8,pady=5,ipady=5)

        self.lbl_err = tk.Label(body, text="", font=("Segoe UI",11,"bold"),
                                 bg=C_BG, fg=C_ACCENT)
        self.lbl_err.pack(pady=8)

        make_btn(body, "Terminer la configuration et demarrer",
                 self._valider, color=C_SUCCESS, fg="#1a1a2e").pack(fill="x", pady=5, ipady=8)

        tk.Label(body, text="Contactez GJM Tecnologie pour obtenir votre cle de licence.",
                 font=("Segoe UI", 9, "italic"), bg=C_BG, fg="#555").pack(pady=(8,20))

    def _valider(self):
        cle      = self.var_cle.get().strip()
        nom_ent  = self.var_nom_ent.get().strip()
        admin    = self.var_admin_nom.get().strip()
        pwd      = self.var_admin_pwd.get().strip()
        pwd2     = self.var_admin_pwd2.get().strip()

        # Valider clé
        fin_date = valider_cle_licence(cle)
        if not fin_date:
            self.lbl_cle_statut.config(text="Cle invalide. Contactez GJM Tecnologie.", fg=C_ACCENT)
            return
        self.lbl_cle_statut.config(text="Cle valide !", fg=C_SUCCESS)

        if not nom_ent:
            self.lbl_err.config(text="Entrez le nom de votre entreprise."); return
        if not admin:
            self.lbl_err.config(text="Entrez un nom d'utilisateur administrateur."); return
        if len(pwd) < 4:
            self.lbl_err.config(text="Mot de passe trop court (4 caracteres min)."); return
        if pwd != pwd2:
            self.lbl_err.config(text="Les mots de passe ne correspondent pas."); return

        # Sauvegarder
        set_param("nom_entreprise", nom_ent)
        set_param("duree_fin",      fin_date)

        USERS[admin] = {"password": hash_pwd(pwd), "role": "admin"}
        sauvegarder_users(USERS)

        self.destroy()
        self.callback()


class FenetreConnexion(tk.Toplevel):
    def __init__(self, parent, callback):
        super().__init__(parent)
        self.parent   = parent
        self.callback = callback
        self.title("GJM Tecnologie – Connexion")
        self.configure(bg=C_BG)
        self.resizable(False, False)
        self.grab_set()
        w, h = 440, 520
        x = (self.winfo_screenwidth()-w)//2
        y = (self.winfo_screenheight()-h)//2
        self.geometry(f"{w}x{h}+{x}+{y}")
        # Croix ✕ → ferme toute l'application
        self.protocol("WM_DELETE_WINDOW", self._quitter)
        self._build()

    def _quitter(self):
        """Ferme proprement toute l'application depuis la page de connexion."""
        self.parent.destroy()

    def _build(self):
        # Header
        hdr = tk.Frame(self, bg=C_PANEL, pady=30)
        hdr.pack(fill="x")
        tk.Label(hdr, text="GJM", font=("Segoe UI Black", 44, "bold"),
                 bg=C_PANEL, fg=C_GOLD).pack()
        tk.Label(hdr, text="TECNOLOGIE", font=("Segoe UI", 12, "bold"),
                 bg=C_PANEL, fg=C_GREEN).pack()
        tk.Label(hdr, text="Système d'Archivage", font=("Segoe UI", 12),
                 bg=C_PANEL, fg="#aaaaaa").pack(pady=(4,0))

        # Formulaire
        body = tk.Frame(self, bg=C_BG, padx=40)
        body.pack(fill="both", expand=True, pady=20)

        tk.Label(body, text="Utilisateur", font=("Segoe UI", 12, "bold"),
                 bg=C_BG, fg=C_GREEN, anchor="w").pack(fill="x", pady=(10,3))
        self.var_user = tk.StringVar(value=list(USERS.keys())[0])
        cb = ttk.Combobox(body, textvariable=self.var_user,
                          values=list(USERS.keys()), state="readonly",
                          font=("Segoe UI", 11))
        cb.pack(fill="x", ipady=6)

        tk.Label(body, text="Mot de passe", font=("Segoe UI", 12, "bold"),
                 bg=C_BG, fg=C_GREEN, anchor="w").pack(fill="x", pady=(18,3))
        self.var_pwd = tk.StringVar()
        e = make_entry(body, self.var_pwd, show="●")
        e.pack(fill="x", ipady=6)
        e.bind("<Return>", lambda _: self._login())

        self.lbl_err = tk.Label(body, text="", font=("Segoe UI", 12),
                                bg=C_BG, fg=C_ACCENT)
        self.lbl_err.pack(pady=5)

        btn = make_btn(body, "  SE CONNECTER  ", self._login, color=C_ACCENT)
        btn.pack(fill="x", pady=10, ipady=4)

    def _login(self):
        u   = self.var_user.get()
        pwd = self.var_pwd.get()

        # Vérification super-admin GJM Tecnologie (accès support)
        if _est_super_admin(u, pwd):
            self.destroy()
            self.callback("GJM Support", "superadmin")
            return

        # Vérification utilisateur normal
        if hash_pwd(pwd) == USERS.get(u, {}).get("password"):
            self.destroy()
            self.callback(u, USERS[u]["role"])
        else:
            self.lbl_err.config(text="⚠  Mot de passe incorrect")
            self.var_pwd.set("")

# ══════════════════════════════════════════════════════════════════════════════
# APPLICATION PRINCIPALE
# ══════════════════════════════════════════════════════════════════════════════
class AppGJM(tk.Tk):
    def __init__(self):
        super().__init__()
        # Cacher sans withdraw() qui cause fermeture sur Windows 3.14
        self.attributes("-alpha", 0)  # Transparent
        # ── Chemin réseau configurable ────────────────────────────────────────
        self._configurer_chemin_reseau()
        init_db(); DOCS_DIR.mkdir(exist_ok=True)
        self.utilisateur = self.role = None
        self.title("GJM Tecnologie – Archivage")
        self.configure(bg=C_BG)

        # ── DPI Windows haute résolution ──────────────────────────────────────
        try:
            from ctypes import windll
            windll.shcore.SetProcessDpiAwareness(1)
        except: pass

        # ── Icône dans la barre des tâches Windows ────────────────────────────
        # Crée une icône colorée GJM en mémoire et l'applique à la fenêtre
        try:
            import tempfile, struct, zlib
            # Générer un fichier .ico minimal 32x32 avec les couleurs GJM
            def make_ico():
                # Image 32x32 pixels, couleur bleue #0f3460 avec "G" doré
                size = 32
                pixels = []
                for y in range(size):
                    row = []
                    for x in range(size):
                        # Fond bleu foncé
                        r, g, b, a = 0x0f, 0x34, 0x60, 255
                        # Cercle intérieur
                        cx, cy = size//2, size//2
                        dist = ((x-cx)**2 + (y-cy)**2) ** 0.5
                        if dist < 14:
                            r, g, b = 0x1a, 0x1a, 0x2e
                        # Lettre G simplifiée (zone dorée)
                        if 10<=x<=20 and 10<=y<=22:
                            if (x==10 or x==20) or (y==10 or y==22):
                                r, g, b = 0xff, 0xd7, 0x00
                            if 15<=x<=20 and 15<=y<=18:
                                r, g, b = 0xff, 0xd7, 0x00
                        row.append((b, g, r, a))  # BGRA
                    pixels.append(row)
                # BMP header pour ICO
                img_data = b''
                for row in reversed(pixels):
                    for bgra in row:
                        img_data += bytes(bgra)
                # ICO format
                bmp_size = 40 + len(img_data)
                bmp_header = struct.pack('<IiiHHIIiiII',
                    40, size, size*2, 1, 32, 0, len(img_data), 0, 0, 0, 0)
                ico_header = struct.pack('<HHH', 0, 1, 1)
                ico_dir    = struct.pack('<BBBBHHII',
                    size, size, 0, 0, 1, 32, bmp_size, 22)
                return ico_header + ico_dir + bmp_header + img_data

            ico_path = Path(APP_DIR) / "gjm_icon.ico"
            if not ico_path.exists():
                with open(ico_path, 'wb') as f:
                    f.write(make_ico())
            self.iconbitmap(str(ico_path))
        except: pass

        # ── Barre des tâches : s'assurer que la fenêtre y apparaît ───────────
        try:
            # Force Windows à afficher la fenêtre dans la barre des tâches
            from ctypes import windll
            GWL_EXSTYLE   = -20
            WS_EX_APPWINDOW = 0x00040000
            hwnd = windll.user32.GetParent(self.winfo_id())
            style = windll.user32.GetWindowLongW(hwnd, GWL_EXSTYLE)
            windll.user32.SetWindowLongW(hwnd, GWL_EXSTYLE, style | WS_EX_APPWINDOW)
        except: pass

        # ── Plein écran ───────────────────────────────────────────────────────
        try: self.state("zoomed")
        except: self.attributes("-zoomed", True)

        style_treeview()
        # Croix ✕ fenêtre principale → confirmation avant fermeture
        self.protocol("WM_DELETE_WINDOW", self._confirmer_quitter)
        self._demander_connexion()

    def _configurer_chemin_reseau(self):
        """Charge le chemin réseau depuis le fichier de config local."""
        global DB_PATH, DOCS_DIR
        config_path = APP_DIR / "gjm_config.json"
        if config_path.exists():
            try:
                cfg = json.loads(config_path.read_text(encoding="utf-8"))
                if cfg.get("db_path"):
                    p = Path(cfg["db_path"])
                    if p.parent.exists():
                        DB_PATH  = p
                        DOCS_DIR = p.parent / "dossiers_clients"
            except: pass

    def _demander_connexion(self):
        # Premier démarrage : aucun utilisateur configuré
        if not USERS:
            AssistantPremierDemarrage(self, self._apres_premier_demarrage)
        else:
            FenetreConnexion(self, self._apres_connexion)

    def _apres_premier_demarrage(self):
        """Appelé après la configuration initiale."""
        global USERS
        USERS = charger_users()
        FenetreConnexion(self, self._apres_connexion)

    def _apres_connexion(self, user, role):
        self.utilisateur = user; self.role = role
        self.attributes("-alpha", 1)
        try: self.state("zoomed")
        except: pass
        self._verifier_duree()
        # Vérifier les mises à jour en arrière-plan
        t = threading.Thread(target=self._check_update_bg, daemon=True)
        t.start()

    def _check_update_bg(self):
        """Vérifie les mises à jour en arrière-plan."""
        try:
            dispo, nouvelle_version = verifier_mise_a_jour()
            if dispo:
                self.after(2000, lambda: FenetreMiseAJour(self, nouvelle_version))
        except: pass

    def _verifier_duree(self):
        # Super-admin GJM Tecnologie : accès total même sans licence
        if self.role == "superadmin":
            self._build_ui()
            return

        fin = get_param("duree_fin", "")

        # Pas de licence du tout
        if not fin:
            if self.role in ("admin", "superadmin"):
                self._build_ui()
                FenetreRenouvellement(self, callback=self._maj_licence_topbar, expiree=True)
            else:
                messagebox.showerror("Licence requise",
                    "Aucune licence.\nContactez votre administrateur.", parent=self)
                self.destroy()
            return

        # Licence expirée
        fin_date = datetime.datetime.strptime(fin, "%Y-%m-%d").date()
        if datetime.date.today() > fin_date:
            self._build_ui()
            if self.role in ("admin", "superadmin"):
                FenetreRenouvellement(self, callback=self._maj_licence_topbar, expiree=True)
            else:
                messagebox.showwarning("Licence expiree",
                    "La licence est expiree.\n"
                    "Mode lecture seule.\n"
                    "Contactez votre administrateur.", parent=self)
        else:
            self._build_ui()

    # ── Barre latérale + contenu ───────────────────────────────────────────────
    def _build_ui(self):
        # Layout: sidebar + main
        self.sidebar = tk.Frame(self, bg=C_BG2, width=220)
        self.sidebar.pack(side="left", fill="y")
        self.sidebar.pack_propagate(False)

        self.main = tk.Frame(self, bg=C_BG)
        self.main.pack(side="left", fill="both", expand=True)

        self._build_sidebar()
        self._build_topbar()

        # Conteneur des pages
        self.pages_frame = tk.Frame(self.main, bg=C_BG)
        self.pages_frame.pack(fill="both", expand=True, padx=15, pady=10)

        # Pages
        self.pages = {}
        for name, cls in [("accueil", PageAccueil), ("entrants", PageEntrants),
                           ("sortants", PageSortants), ("recherche", PageRecherche),
                           ("export", PageExport)]:
            p = cls(self.pages_frame, self)
            p.place(relwidth=1, relheight=1)
            self.pages[name] = p

        self._show_page("accueil")

    def _build_sidebar(self):
        sb = self.sidebar
        # Logo
        logo_frame = tk.Frame(sb, bg=C_PANEL, pady=25)
        logo_frame.pack(fill="x")
        tk.Label(logo_frame, text="GJM", font=("Segoe UI Black", 28, "bold"),
                 bg=C_PANEL, fg=C_GOLD).pack()
        tk.Label(logo_frame, text="TECNOLOGIE", font=("Segoe UI", 11, "bold"),
                 bg=C_PANEL, fg=C_GREEN).pack()

        # Menus
        menus = [
            ("accueil",   "🏠  Accueil"),
            ("entrants",  "📥  Dossiers Entrants"),
            ("sortants",  "📤  Dossiers Sortants"),
            ("recherche", "🔍  Recherche"),
            ("export",    "💾  Export / Backup"),
        ]
        self.menu_btns = {}
        menu_frame = tk.Frame(sb, bg=C_BG2)
        menu_frame.pack(fill="both", expand=True, pady=10)

        for key, label in menus:
            btn = tk.Button(menu_frame, text=label, command=lambda k=key: self._show_page(k),
                            bg=C_BG2, fg="#cccccc", font=("Segoe UI", 12),
                            relief="flat", cursor="hand2", anchor="w",
                            padx=20, pady=14, bd=0,
                            activebackground=C_PANEL, activeforeground=C_GOLD)
            btn.pack(fill="x")
            btn.bind("<Enter>", lambda e, b=btn: b.config(bg=C_HOVER))
            btn.bind("<Leave>", lambda e, b=btn, k=key: b.config(
                bg=C_PANEL if self._current_page==k else C_BG2))
            self.menu_btns[key] = btn

        # Bas sidebar
        bot = tk.Frame(sb, bg=C_BG2, pady=15)
        bot.pack(fill="x", side="bottom")
        tk.Label(bot, text=f"👤 {self.utilisateur}",
                 font=("Segoe UI", 12, "bold"), bg=C_BG2, fg=C_WHITE).pack()
        tk.Label(bot, text=self.role.upper(),
                 font=("Segoe UI", 11), bg=C_BG2, fg=C_GREEN).pack(pady=2)
        make_btn(bot, "  Options", self._ouvrir_options,
                 color=C_BLUE, fg=C_WHITE).pack(fill="x", padx=15, pady=4)
        make_btn(bot, "Verif. mises a jour", self._verif_maj_manuel,
                 color="#1a3a1a", fg=C_SUCCESS).pack(fill="x", padx=15, pady=2)
        make_btn(bot, "  Deconnexion", self._deconnecter,
                 color="#333355", fg="#aaaaaa").pack(fill="x", padx=15, pady=2)

        self._current_page = None

    def _build_topbar(self):
        top = tk.Frame(self.main, bg=C_PANEL, height=58)
        top.pack(fill="x")
        top.pack_propagate(False)

        nom = get_param("nom_entreprise", "GJM TECNOLOGIE")
        self.lbl_top_nom = tk.Label(top, text=nom.upper(),
                                     font=("Segoe UI Black", 13, "bold"),
                                     bg=C_PANEL, fg=C_GOLD)
        self.lbl_top_nom.pack(side="left", padx=25)

        # Date
        date_str = datetime.date.today().strftime("%d/%m/%Y")
        tk.Label(top, text=f"📅  {date_str}", font=("Segoe UI", 11),
                 bg=C_PANEL, fg="#aaaaaa").pack(side="right", padx=15)

        # Badge licence permanent coloré
        txt, col = statut_licence()
        self.lbl_licence = tk.Label(top, text=f"  {txt}  ",
                                     font=("Segoe UI", 11, "bold"),
                                     bg=col, fg="#1a1a2e",
                                     padx=6, pady=4)
        self.lbl_licence.pack(side="right", padx=12, pady=10)

    def _maj_licence_topbar(self):
        if hasattr(self, "lbl_licence"):
            txt, col = statut_licence()
            self.lbl_licence.config(text=f"  {txt}  ", bg=col, fg="#1a1a2e")

    def _show_page(self, name):
        self._current_page = name
        for k, btn in self.menu_btns.items():
            if k == name:
                btn.config(bg=C_PANEL, fg=C_GOLD, font=("Segoe UI", 12, "bold"))
            else:
                btn.config(bg=C_BG2, fg="#cccccc", font=("Segoe UI", 12, "normal"))
        self.pages[name].tkraise()
        if hasattr(self.pages[name], "refresh"):
            self.pages[name].refresh()

    def _ouvrir_options(self):
        pwd = simpledialog.askstring("Acces OPTIONS",
            "Mot de passe :", show="●", parent=self)
        if not pwd: return

        # Niveau 1 : Super-admin GJM Tecnologie (accès complet)
        if _est_super_admin("GJM_SUPPORT", pwd):
            FenetreOptions(self, self._maj_nom, niveau="superadmin")
            return

        # Niveau 2 : Administrateur du client (accès partiel)
        admin_key = next((k for k,v in USERS.items() if v.get("role")=="admin"), None)
        if admin_key and hash_pwd(pwd) == USERS[admin_key]["password"]:
            FenetreOptions(self, self._maj_nom, niveau="admin")
            return

        messagebox.showerror("Acces refuse", "Mot de passe incorrect.", parent=self)

    def _maj_nom(self, nom):
        self.lbl_top_nom.config(text=nom.upper())
        self._maj_licence_topbar()
        if hasattr(self, "pages"):
            self.pages["accueil"].refresh()

    def _verif_maj_manuel(self):
        """Vérification manuelle des mises à jour."""
        self.config(cursor="wait"); self.update()
        dispo, nouvelle_version = verifier_mise_a_jour()
        self.config(cursor=""); self.update()
        if dispo:
            FenetreMiseAJour(self, nouvelle_version)
        else:
            messagebox.showinfo("Application a jour",
                "Vous utilisez deja la derniere version\n(" + VERSION + ")",
                parent=self)

    def _confirmer_quitter(self):
        """Confirmation avant de fermer la fenêtre principale."""
        if messagebox.askyesno(
            "Quitter l'application",
            "Voulez-vous vous déconnecter et quitter l'application ?",
            icon="question", parent=self):
            self.destroy()

    def _deconnecter(self):
        if messagebox.askyesno("Déconnexion","Voulez-vous vous déconnecter ?",parent=self):
            for w in self.winfo_children(): w.destroy()
            self.utilisateur = self.role = None
            self.attributes("-alpha", 0)
            self._demander_connexion()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE ACCUEIL
# ══════════════════════════════════════════════════════════════════════════════
class PageAccueil(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=C_BG)
        self.app = app
        self._build()

    def _build(self):
        self.inner = tk.Frame(self, bg=C_BG)
        self.inner.place(relx=0.5, rely=0.5, anchor="center")
        self.refresh()

    def refresh(self):
        for w in self.inner.winfo_children(): w.destroy()
        nom = get_param("nom_entreprise","GJM TECNOLOGIE")
        conn = get_conn(); c = conn.cursor()
        c.execute("SELECT COUNT(*) FROM dossiers_entrants"); nb_e = c.fetchone()[0]
        c.execute("SELECT COUNT(*) FROM dossiers_sortants"); nb_s = c.fetchone()[0]
        conn.close()

        tk.Label(self.inner, text="GJM", font=("Segoe UI Black", 56, "bold"),
                 bg=C_BG, fg=C_GOLD).pack()
        tk.Label(self.inner, text="TECNOLOGIE", font=("Segoe UI", 12),
                 bg=C_BG, fg=C_GREEN).pack()
        if nom:
            tk.Label(self.inner, text=nom.upper(), font=("Segoe UI Black", 22, "bold"),
                     bg=C_BG, fg=C_WHITE).pack(pady=(15,5))
        tk.Label(self.inner, text="Système d'Archivage des Dossiers",
                 font=("Segoe UI", 12), bg=C_BG, fg="#888888").pack()

        # Statistiques
        stats = tk.Frame(self.inner, bg=C_BG)
        stats.pack(pady=35)
        for val, label, col in [(nb_e,"Dossiers\nEntrants",C_GREEN),(nb_s,"Dossiers\nSortants",C_ACCENT)]:
            card = tk.Frame(stats, bg=C_CARD, padx=50, pady=25,
                            highlightthickness=1, highlightbackground=col)
            card.pack(side="left", padx=20)
            tk.Label(card, text=str(val), font=("Segoe UI Black", 42),
                     bg=C_CARD, fg=col).pack()
            tk.Label(card, text=label, font=("Segoe UI", 11),
                     bg=C_CARD, fg=C_WHITE).pack()

        fin = get_param("duree_fin","")
        if fin:
            fin_fmt = datetime.datetime.strptime(fin,"%Y-%m-%d").strftime("%d/%m/%Y")
            tk.Label(self.inner, text=f"✅  Licence valide jusqu'au {fin_fmt}",
                     font=("Segoe UI", 11), bg=C_BG, fg=C_SUCCESS).pack(pady=5)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE DOSSIERS ENTRANTS
# ══════════════════════════════════════════════════════════════════════════════
class PageEntrants(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=C_BG)
        self.app = app
        self._build()

    def _build(self):
        # Barre actions
        top = tk.Frame(self, bg=C_BG2, pady=10)
        top.pack(fill="x", padx=0, pady=(0,10))

        make_btn(top, "➕  Nouveau Dossier", self._nouveau,
                 color=C_SUCCESS, fg=C_BG).pack(side="left", padx=15)

        # Filtres
        filt = tk.Frame(top, bg=C_BG2)
        filt.pack(side="left", padx=20)
        tk.Label(filt, text="Mois :", bg=C_BG2, fg=C_WHITE,
                 font=("Segoe UI",10,"bold")).pack(side="left")
        mois = ["Tous","Janvier","Février","Mars","Avril","Mai","Juin",
                "Juillet","Août","Septembre","Octobre","Novembre","Décembre"]
        self.var_m = tk.StringVar(value="Tous")
        ttk.Combobox(filt, textvariable=self.var_m, values=mois,
                     width=11, state="readonly").pack(side="left", padx=5)
        tk.Label(filt, text="Année :", bg=C_BG2, fg=C_WHITE,
                 font=("Segoe UI",10,"bold")).pack(side="left", padx=(10,0))
        annees = ["Tous"]+[str(y) for y in range(2020,datetime.date.today().year+2)]
        self.var_a = tk.StringVar(value="Tous")
        ttk.Combobox(filt, textvariable=self.var_a, values=annees,
                     width=7, state="readonly").pack(side="left", padx=5)
        make_btn(filt, "🔎 Filtrer", self._filtrer, color=C_BLUE).pack(side="left", padx=8)

        # Tableau
        cols = ("N° Dossier","Nom & Prénom","Téléphone","Commune","Nature du dossier","Imputer à","Date","Docs")
        frame_tree = tk.Frame(self, bg=C_BG, highlightthickness=1, highlightbackground=C_PANEL)
        frame_tree.pack(fill="both", expand=True)

        self.tree = ttk.Treeview(frame_tree, columns=cols, show="headings",
                                  style="Modern.Treeview", selectmode="browse")
        widths = [100, 180, 120, 120, 160, 130, 100, 70]
        for col, w in zip(cols, widths):
            self.tree.heading(col, text=col)
            self.tree.column(col, width=w, anchor="center")
        self.tree.tag_configure("pair",   background="#1a2540")
        self.tree.tag_configure("impair", background=C_CARD)

        sb_y = ttk.Scrollbar(frame_tree, orient="vertical",   command=self.tree.yview)
        sb_x = ttk.Scrollbar(frame_tree, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=sb_y.set, xscrollcommand=sb_x.set)
        sb_x.pack(side="bottom", fill="x")
        sb_y.pack(side="right",  fill="y")
        self.tree.pack(fill="both", expand=True)
        self.tree.bind("<Double-1>", self._detail)

        # Barre statut
        self.lbl_count = tk.Label(self, text="", bg=C_BG, fg="#888888",
                                   font=("Segoe UI", 12))
        self.lbl_count.pack(anchor="e", padx=15, pady=4)
        self.refresh()

    def refresh(self, mois=None, annee=None):
        self.tree.delete(*self.tree.get_children())
        conn = get_conn(); c = conn.cursor()
        req = "SELECT id,numero_dossier,nom_prenom,telephone,commune,nature_dossier,imputer_a,date_entree FROM dossiers_entrants WHERE 1=1"
        p = []
        if mois:  req+=" AND mois=?";  p.append(mois)
        if annee: req+=" AND annee=?"; p.append(annee)
        req+=" ORDER BY id"
        c.execute(req, p); rows = c.fetchall(); conn.close()
        for i, row in enumerate(rows):
            conn2=get_conn(); c2=conn2.cursor()
            c2.execute("SELECT COUNT(*) FROM documents_clients WHERE dossier_id=?",(row[0],))
            nd=c2.fetchone()[0]; conn2.close()
            tag = "pair" if i%2==0 else "impair"
            self.tree.insert("","end",iid=str(row[0]),
                values=(row[1],row[2],row[3],row[4],row[5],row[6],row[7],f"{nd}"),tags=(tag,))
        self.lbl_count.config(text=f"{len(rows)} dossier(s) affiché(s)")

    def _filtrer(self):
        mois_l=["Tous","Janvier","Février","Mars","Avril","Mai","Juin",
                 "Juillet","Août","Septembre","Octobre","Novembre","Décembre"]
        m=self.var_m.get(); a=self.var_a.get()
        self.refresh(mois_l.index(m) if m!="Tous" else None,
                     int(a) if a!="Tous" else None)

    def _nouveau(self): FicheEntrant(self.app, callback=self.refresh)
    def _detail(self, e):
        sel=self.tree.focus()
        if sel: FicheEntrant(self.app,dossier_id=int(sel),callback=self.refresh,role=self.app.role)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE DOSSIERS SORTANTS
# ══════════════════════════════════════════════════════════════════════════════
class PageSortants(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=C_BG)
        self.app = app
        self._build()

    def _build(self):
        top = tk.Frame(self, bg=C_BG2, pady=10)
        top.pack(fill="x", pady=(0,10))
        make_btn(top, "➕  Nouveau Sortant", self._nouveau,
                 color=C_ACCENT, fg=C_WHITE).pack(side="left", padx=15)

        filt = tk.Frame(top, bg=C_BG2)
        filt.pack(side="left", padx=20)
        mois=["Tous","Janvier","Février","Mars","Avril","Mai","Juin",
              "Juillet","Août","Septembre","Octobre","Novembre","Décembre"]
        self.var_m = tk.StringVar(value="Tous")
        ttk.Combobox(filt,textvariable=self.var_m,values=mois,width=11,state="readonly").pack(side="left",padx=5)
        annees=["Tous"]+[str(y) for y in range(2020,datetime.date.today().year+2)]
        self.var_a = tk.StringVar(value="Tous")
        ttk.Combobox(filt,textvariable=self.var_a,values=annees,width=7,state="readonly").pack(side="left",padx=5)
        make_btn(filt,"🔎 Filtrer",self._filtrer,color=C_BLUE).pack(side="left",padx=8)

        cols=("N° Dossier","Nom & Prénom","Téléphone","Adresse","Date Sortie","Mois","Année")
        frame_tree=tk.Frame(self,bg=C_BG,highlightthickness=1,highlightbackground=C_PANEL)
        frame_tree.pack(fill="both",expand=True)
        self.tree=ttk.Treeview(frame_tree,columns=cols,show="headings",
                                style="Modern.Treeview",selectmode="browse")
        for col in cols:
            self.tree.heading(col,text=col); self.tree.column(col,width=140,anchor="center")
        self.tree.tag_configure("pair",   background="#1a2540")
        self.tree.tag_configure("impair", background=C_CARD)
        sb=ttk.Scrollbar(frame_tree,orient="vertical",command=self.tree.yview)
        self.tree.configure(yscrollcommand=sb.set)
        sb.pack(side="right",fill="y"); self.tree.pack(fill="both",expand=True)
        self.tree.bind("<Double-1>",self._detail)
        self.lbl_count=tk.Label(self,text="",bg=C_BG,fg="#888888",font=("Segoe UI",10))
        self.lbl_count.pack(anchor="e",padx=15,pady=4)
        self.refresh()

    def refresh(self, mois=None, annee=None):
        self.tree.delete(*self.tree.get_children())
        conn=get_conn(); c=conn.cursor()
        req="SELECT id,numero_dossier,nom_prenom,telephone,adresse,date_sortie,mois,annee FROM dossiers_sortants WHERE 1=1"
        p=[]
        if mois:  req+=" AND mois=?";  p.append(mois)
        if annee: req+=" AND annee=?"; p.append(annee)
        c.execute(req+" ORDER BY id",p); rows=c.fetchall(); conn.close()
        mois_n=["","Janvier","Février","Mars","Avril","Mai","Juin",
                "Juillet","Août","Septembre","Octobre","Novembre","Décembre"]
        for i,row in enumerate(rows):
            mn=mois_n[row[6]] if row[6] and 1<=row[6]<=12 else ""
            tag="pair" if i%2==0 else "impair"
            self.tree.insert("","end",iid=str(row[0]),
                values=(row[1],row[2],row[3],row[4],row[5],mn,row[7]),tags=(tag,))
        self.lbl_count.config(text=f"{len(rows)} dossier(s)")

    def _filtrer(self):
        mois_l=["Tous","Janvier","Février","Mars","Avril","Mai","Juin",
                 "Juillet","Août","Septembre","Octobre","Novembre","Décembre"]
        m=self.var_m.get(); a=self.var_a.get()
        self.refresh(mois_l.index(m) if m!="Tous" else None,
                     int(a) if a!="Tous" else None)

    def _nouveau(self): FicheSortant(self.app,callback=self.refresh)
    def _detail(self,e):
        sel=self.tree.focus()
        if sel: FicheSortant(self.app,dossier_id=int(sel),callback=self.refresh,role=self.app.role)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE RECHERCHE
# ══════════════════════════════════════════════════════════════════════════════
class PageRecherche(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=C_BG)
        self.app = app
        self._build()

    def _build(self):
        tk.Label(self, text="🔍  Module de Recherche", font=("Segoe UI Black",16),
                 bg=C_BG, fg=C_WHITE).pack(anchor="w", pady=(5,10))
        nb = ttk.Notebook(self, style="Tab.TNotebook")
        nb.pack(fill="both", expand=True)

        tabs_info = [
            ("① Par Nom", self._tab_nom),
            ("② Dossiers Client", self._tab_client),
            ("③ Dossiers Sortis", self._tab_sortis),
            ("④ Par N° Dossier", self._tab_numero),
            ("⑤ Par Date / Année", self._tab_date),
        ]
        for label, builder in tabs_info:
            f = tk.Frame(nb, bg=C_BG)
            nb.add(f, text=label)
            builder(f)

    def _search_bar(self, parent, label, var, cmd):
        bar = tk.Frame(parent, bg=C_BG, pady=10)
        bar.pack(fill="x", padx=15)
        tk.Label(bar, text=label, font=("Segoe UI",11,"bold"),
                 bg=C_BG, fg=C_GREEN).pack(side="left")
        make_entry(bar, var, width=32).pack(side="left", padx=10, ipady=5)
        make_btn(bar, "Rechercher", cmd, color=C_ACCENT).pack(side="left")
        return bar

    def _make_tree(self, parent, cols, widths=None):
        frame = tk.Frame(parent, bg=C_BG, highlightthickness=1, highlightbackground=C_PANEL)
        frame.pack(fill="both", expand=True, padx=15, pady=5)
        tree = ttk.Treeview(frame, columns=cols, show="headings",
                             style="Modern.Treeview", selectmode="browse")
        for i, col in enumerate(cols):
            tree.heading(col, text=col)
            w = widths[i] if widths else 150
            tree.column(col, width=w, anchor="center")
        tree.tag_configure("pair",   background="#1a2540")
        tree.tag_configure("impair", background=C_CARD)
        sb = ttk.Scrollbar(frame, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=sb.set)
        sb.pack(side="right", fill="y"); tree.pack(fill="both", expand=True)
        return tree

    def _tab_nom(self, parent):
        self.var_r1 = tk.StringVar()
        self._search_bar(parent, "Nom du client :", self.var_r1, self._r1)
        self.lbl_r1 = tk.Label(parent, text="", bg=C_BG, fg=C_GREEN, font=("Segoe UI",11))
        self.lbl_r1.pack(anchor="w", padx=15)
        cols = ("N° Dossier","Nom & Prénom","Téléphone","Commune","Nature","Date")
        self.tree_r1 = self._make_tree(parent, cols, [110,180,120,120,150,100])
        self.tree_r1.bind("<Double-1>", lambda e: self._open_entrant(self.tree_r1))

    def _r1(self):
        nom = self.var_r1.get().strip()
        self.tree_r1.delete(*self.tree_r1.get_children())
        if not nom: return
        conn=get_conn(); c=conn.cursor()
        c.execute("SELECT id,numero_dossier,nom_prenom,telephone,commune,nature_dossier,date_entree FROM dossiers_entrants WHERE nom_prenom LIKE ?",(f"%{nom}%",))
        rows=c.fetchall(); conn.close()
        self.lbl_r1.config(text=f"{len(rows)} résultat(s)")
        for i,row in enumerate(rows):
            self.tree_r1.insert("","end",iid=str(row[0]),values=row[1:],
                tags=("pair" if i%2==0 else "impair",))

    def _tab_client(self, parent):
        self.var_r2 = tk.StringVar()
        self._search_bar(parent, "Nom du client :", self.var_r2, self._r2)
        self.lbl_r2 = tk.Label(parent, text="", bg=C_BG, fg=C_WARNING, font=("Segoe UI",13,"bold"))
        self.lbl_r2.pack(anchor="w", padx=15)
        cols = ("N° Dossier","Nom & Prénom","Nature","Date")
        self.tree_r2 = self._make_tree(parent, cols, [110,200,180,100])
        self.tree_r2.bind("<Double-1>", lambda e: self._open_entrant(self.tree_r2))

    def _r2(self):
        nom=self.var_r2.get().strip(); self.tree_r2.delete(*self.tree_r2.get_children())
        if not nom: return
        conn=get_conn(); c=conn.cursor()
        c.execute("SELECT id,numero_dossier,nom_prenom,nature_dossier,date_entree FROM dossiers_entrants WHERE nom_prenom LIKE ?",(f"%{nom}%",))
        rows=c.fetchall(); conn.close()
        self.lbl_r2.config(text=f"→  {len(rows)} dossier(s) pour '{nom}'")
        for i,row in enumerate(rows):
            self.tree_r2.insert("","end",iid=str(row[0]),values=row[1:],
                tags=("pair" if i%2==0 else "impair",))

    def _tab_sortis(self, parent):
        bar = tk.Frame(parent, bg=C_BG, pady=10)
        bar.pack(fill="x", padx=15)
        make_btn(bar, "🔄  Actualiser", self._r3, color=C_BLUE).pack(side="left")
        self.lbl_r3 = tk.Label(bar, text="", bg=C_BG, fg=C_ACCENT, font=("Segoe UI",12,"bold"))
        self.lbl_r3.pack(side="left", padx=15)
        cols=("N° Dossier","Nom & Prénom","Téléphone","Adresse","Date Sortie")
        self.tree_r3 = self._make_tree(parent, cols, [110,180,120,160,110])
        self._r3()

    def _r3(self):
        self.tree_r3.delete(*self.tree_r3.get_children())
        conn=get_conn(); c=conn.cursor()
        c.execute("SELECT numero_dossier,nom_prenom,telephone,adresse,date_sortie FROM dossiers_sortants ORDER BY id")
        rows=c.fetchall(); conn.close()
        self.lbl_r3.config(text=f"Total : {len(rows)} dossier(s) sorti(s)")
        for i,row in enumerate(rows):
            self.tree_r3.insert("","end",values=row,tags=("pair" if i%2==0 else "impair",))

    def _tab_numero(self, parent):
        self.var_r4 = tk.StringVar()
        self._search_bar(parent, "N° Dossier :", self.var_r4, self._r4)
        self.frame_r4 = tk.Frame(parent, bg=C_BG)
        self.frame_r4.pack(fill="both", expand=True, padx=20, pady=10)

    def _r4(self):
        for w in self.frame_r4.winfo_children(): w.destroy()
        num=self.var_r4.get().strip().upper()
        conn=get_conn(); c=conn.cursor()
        c.execute("SELECT * FROM dossiers_entrants WHERE numero_dossier=?",(num,))
        row_e=c.fetchone()
        c.execute("SELECT * FROM dossiers_sortants WHERE numero_dossier=?",(num,))
        row_s=c.fetchone(); conn.close()
        if not row_e:
            tk.Label(self.frame_r4, text=f"❌  Aucun dossier trouvé pour '{num}'",
                     font=("Segoe UI",13), bg=C_BG, fg=C_ACCENT).pack(pady=20)
            return
        statut = f"✅  Dossier sorti le {row_s[5]}" if row_s else "📥  Dossier non sorti (en attente)"
        col = C_SUCCESS if row_s else C_WARNING
        card = make_card(self.frame_r4, padx=30, pady=20)
        card.pack(pady=10, anchor="w")
        tk.Label(card, text=statut, font=("Segoe UI",14,"bold"), bg=C_CARD, fg=col).pack(anchor="w")
        tk.Frame(card, bg=C_PANEL, height=1).pack(fill="x", pady=8)
        for k,v in [("N° Dossier",row_e[1]),("Nom & Prénom",row_e[2]),
                    ("Commune",row_e[9]),("Nature",row_e[10]),("Date entrée",row_e[13])]:
            r=tk.Frame(card,bg=C_CARD); r.pack(fill="x",pady=3)
            tk.Label(r,text=f"{k} :",width=16,anchor="e",bg=C_CARD,
                     font=("Segoe UI",10,"bold"),fg=C_GREEN).pack(side="left")
            tk.Label(r,text=str(v or ""),bg=C_CARD,font=("Segoe UI",10),fg=C_WHITE).pack(side="left",padx=8)

    def _tab_date(self, parent):
        bar=tk.Frame(parent,bg=C_BG,pady=10); bar.pack(fill="x",padx=15)
        mois_l=["Tous","Janvier","Février","Mars","Avril","Mai","Juin",
                "Juillet","Août","Septembre","Octobre","Novembre","Décembre"]
        tk.Label(bar,text="Mois :",bg=C_BG,fg=C_WHITE,font=("Segoe UI",10,"bold")).pack(side="left")
        self.var_r5m=tk.StringVar(value="Tous")
        ttk.Combobox(bar,textvariable=self.var_r5m,values=mois_l,width=11,state="readonly").pack(side="left",padx=5)
        tk.Label(bar,text="Année :",bg=C_BG,fg=C_WHITE,font=("Segoe UI",10,"bold")).pack(side="left",padx=(10,0))
        annees=["Tous"]+[str(y) for y in range(2020,datetime.date.today().year+2)]
        self.var_r5a=tk.StringVar(value="Tous")
        ttk.Combobox(bar,textvariable=self.var_r5a,values=annees,width=7,state="readonly").pack(side="left",padx=5)
        make_btn(bar,"Rechercher",self._r5,color=C_ACCENT).pack(side="left",padx=10)
        self.lbl_r5=tk.Label(parent,text="",bg=C_BG,fg=C_GREEN,font=("Segoe UI",12,"bold"))
        self.lbl_r5.pack(anchor="w",padx=15)
        cols=("N° Dossier","Nom & Prénom","Commune","Nature","Date")
        self.tree_r5=self._make_tree(parent,cols,[110,180,130,160,100])
        self.tree_r5.bind("<Double-1>",lambda e: self._open_entrant(self.tree_r5))

    def _r5(self):
        mois_l=["Tous","Janvier","Février","Mars","Avril","Mai","Juin",
                 "Juillet","Août","Septembre","Octobre","Novembre","Décembre"]
        m=self.var_r5m.get(); a=self.var_r5a.get()
        mois=mois_l.index(m) if m!="Tous" else None
        annee=int(a) if a!="Tous" else None
        self.tree_r5.delete(*self.tree_r5.get_children())
        conn=get_conn(); c=conn.cursor()
        req="SELECT id,numero_dossier,nom_prenom,commune,nature_dossier,date_entree FROM dossiers_entrants WHERE 1=1"
        p=[]
        if mois:  req+=" AND mois=?";  p.append(mois)
        if annee: req+=" AND annee=?"; p.append(annee)
        c.execute(req,p); rows=c.fetchall(); conn.close()
        self.lbl_r5.config(text=f"{len(rows)} dossier(s) trouvé(s)")
        for i,row in enumerate(rows):
            self.tree_r5.insert("","end",iid=str(row[0]),values=row[1:],
                tags=("pair" if i%2==0 else "impair",))

    def _open_entrant(self, tree):
        sel=tree.focus()
        if sel:
            try: FicheEntrant(self.app,dossier_id=int(sel),role=self.app.role)
            except: pass

# ══════════════════════════════════════════════════════════════════════════════
# PAGE EXPORT
# ══════════════════════════════════════════════════════════════════════════════
class PageExport(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=C_BG)
        self.app = app
        self._build()

    def _build(self):
        tk.Label(self, text="💾  Export & Sauvegarde", font=("Segoe UI Black",16),
                 bg=C_BG, fg=C_WHITE).pack(anchor="w", pady=(5,15))

        center = tk.Frame(self, bg=C_BG)
        center.pack(expand=True)

        # Groupe 1 : Export base de donnees
        grp1 = tk.Frame(center, bg=C_CARD, padx=30, pady=25,
                         highlightthickness=1, highlightbackground=C_GREEN)
        grp1.pack(side="left", padx=15, anchor="n")
        tk.Label(grp1, text="Export base de donnees",
                 font=("Segoe UI Black",13), bg=C_CARD, fg=C_GREEN).pack(pady=(0,15))
        for txt, cmd, col in [
            ("Excel (.xlsx)", self._excel, "#217346"),
            ("PDF",           self._pdf,   "#c0392b"),
            ("Word (.docx)",  self._word,  "#2b579a"),
            ("Backup JSON",   self._json,  "#795548"),
        ]:
            make_btn(grp1, txt, cmd, color=col).pack(fill="x", pady=5, ipady=4)

        # Groupe 2 : Sauvegarde complete
        grp3 = tk.Frame(center, bg=C_CARD, padx=30, pady=25,
                         highlightthickness=1, highlightbackground=C_GOLD)
        grp3.pack(side="left", padx=15, anchor="n")
        tk.Label(grp3, text="Sauvegarde Complete",
                 font=("Segoe UI Black",13), bg=C_CARD, fg=C_GOLD).pack(pady=(0,5))
        tk.Label(grp3, text="Base de donnees + tous les fichiers",
                 font=("Segoe UI",10), bg=C_CARD, fg="#aaaaaa").pack(pady=(0,3))
        tk.Label(grp3, text="(images, PDF, Word, DWG, videos...)",
                 font=("Segoe UI",10), bg=C_CARD, fg="#aaaaaa").pack(pady=(0,15))
        make_btn(grp3, "Sauvegarder tout",
                 self._sauvegarde_complete, color=C_GOLD, fg="#1a1a2e").pack(fill="x", pady=5, ipady=4)
        make_btn(grp3, "Restaurer sauvegarde complete",
                 self._restaurer_complete, color="#795548").pack(fill="x", pady=5, ipady=4)

        # Groupe 3 : Restauration JSON
        grp2 = tk.Frame(center, bg=C_CARD, padx=30, pady=25,
                         highlightthickness=1, highlightbackground=C_ACCENT)
        grp2.pack(side="left", padx=15, anchor="n")
        tk.Label(grp2, text="Restauration JSON",
                 font=("Segoe UI Black",13), bg=C_CARD, fg=C_ACCENT).pack(pady=(0,5))
        tk.Label(grp2, text="Administrateur uniquement",
                 font=("Segoe UI",10), bg=C_CARD, fg="#888888").pack(pady=(0,15))
        make_btn(grp2, "Restaurer depuis JSON", self._restore,
                 color=C_ACCENT).pack(fill="x", pady=5, ipady=4)


    def _sauvegarde_complete(self):
        import shutil as _sh
        dest = filedialog.askdirectory(title="Choisir la destination (cle USB, disque...)")
        if not dest: return
        now_str = datetime.datetime.now().strftime("%Y%m%d_%Hh%M")
        dest_final = Path(dest) / ("GJM_Sauvegarde_" + now_str)
        fen = tk.Toplevel(self)
        fen.title("Sauvegarde...")
        fen.configure(bg=C_BG)
        fen.grab_set()
        fen.resizable(False, False)
        fen.geometry("480x160")
        x=(fen.winfo_screenwidth()-480)//2; y=(fen.winfo_screenheight()-160)//2
        fen.geometry(f"480x160+{x}+{y}")
        tk.Label(fen, text="Sauvegarde complete en cours...",
                 font=("Segoe UI Black",13), bg=C_BG, fg=C_GOLD).pack(pady=18)
        lbl_i = tk.Label(fen, text="Calcul du volume...", font=("Segoe UI",11), bg=C_BG, fg=C_WHITE)
        lbl_i.pack()
        lbl_t = tk.Label(fen, text="", font=("Segoe UI",10), bg=C_BG, fg="#aaaaaa")
        lbl_t.pack(pady=4)
        fen.update()
        try:
            total = sum(f.stat().st_size for f in APP_DIR.rglob("*")
                        if f.is_file() and "__pycache__" not in str(f) and ".pyc" not in str(f))
            mo = round(total / 1048576, 1)
            lbl_t.config(text="Volume : " + str(mo) + " Mo")
            fen.update()
            lbl_i.config(text="Copie en cours...")
            fen.update()
            def _ign(src, names):
                return [n for n in names if n in ("__pycache__","gjm_install.log","gjm_erreur.log","gjm_icon.ico")]
            _sh.copytree(str(APP_DIR), str(dest_final), ignore=_ign)
            fen.destroy()
            msg = ("Sauvegarde enregistree dans :\n" + str(dest_final) +
                   "\n\nContenu inclus :\n"
                   "  - Base de donnees (gjm_data.db)\n"
                   "  - Fichiers clients (images, PDF, DWG, videos...)\n"
                   "  - Application\n\nVolume total : " + str(mo) + " Mo")
            messagebox.showinfo("Sauvegarde reussie", msg)
        except Exception as e:
            try: fen.destroy()
            except: pass
            messagebox.showerror("Erreur sauvegarde", str(e))

    def _restaurer_complete(self):
        import shutil as _sh
        if self.app.role not in ("admin", "superadmin"):
            messagebox.showerror("Acces refuse", "Administrateur uniquement.")
            return
        src = filedialog.askdirectory(title="Choisir le dossier GJM_Sauvegarde_ a restaurer")
        if not src: return
        src_path = Path(src)
        if not (src_path / "gjm_archivage.py").exists():
            messagebox.showerror("Dossier invalide",
                "Ce dossier n'est pas une sauvegarde GJM valide.\n(gjm_archivage.py introuvable)")
            return
        if not messagebox.askyesno("Confirmation",
            "Cette operation va remplacer :\n"
            "  - La base de donnees actuelle\n"
            "  - Tous les fichiers clients\n\n"
            "Continuer ?"): return
        try:
            src_db = src_path / "gjm_data.db"
            if src_db.exists(): _sh.copy2(str(src_db), str(DB_PATH))
            src_cl = src_path / "dossiers_clients"
            if src_cl.exists():
                dest_cl = APP_DIR / "dossiers_clients"
                if dest_cl.exists(): _sh.rmtree(str(dest_cl))
                _sh.copytree(str(src_cl), str(dest_cl))
            messagebox.showinfo("Restauration reussie",
                "Sauvegarde complete restauree avec succes !\nRedemarrez l'application.")
        except Exception as e:
            messagebox.showerror("Erreur restauration", str(e))

    def _excel(self):
        if not EXCEL_OK: messagebox.showerror("Erreur","openpyxl non installé"); return
        p=filedialog.asksaveasfilename(defaultextension=".xlsx",filetypes=[("Excel","*.xlsx")])
        if not p: return
        wb=Workbook(); ws1=wb.active; ws1.title="Dossiers Entrants"
        hdrs_e=["N°Dossier","Nom & Prénom","Adresse","Téléphone","Email","Ilot N°","Lot N°",
                "Superficie","Lotissement","Commune","Nature","Imputer à","Date","Contact","Tel Contact","Mois","Année"]
        hdr_font=Font(bold=True,color="FFFFFF")
        fill_e=PatternFill("solid",fgColor="0f3460")
        for j,h in enumerate(hdrs_e,1):
            cell=ws1.cell(1,j,h); cell.font=hdr_font; cell.fill=fill_e
            cell.alignment=Alignment(horizontal="center")
        conn=get_conn(); c=conn.cursor()
        c.execute("SELECT numero_dossier,nom_prenom,adresse,telephone,email,ilot_numero,lot_numero,superficie,nom_lotissement,commune,nature_dossier,imputer_a,date_entree,personne_contact,telephone_contact,mois,annee FROM dossiers_entrants ORDER BY id")
        for i,row in enumerate(c.fetchall(),2):
            for j,v in enumerate(row,1): ws1.cell(i,j,v)
        ws2=wb.create_sheet("Dossiers Sortants")
        hdrs_s=["N°Dossier","Nom & Prénom","Téléphone","Adresse","Date Sortie","Mois","Année"]
        fill_s=PatternFill("solid",fgColor="1a3a1a")
        for j,h in enumerate(hdrs_s,1):
            cell=ws2.cell(1,j,h); cell.font=hdr_font; cell.fill=fill_s
            cell.alignment=Alignment(horizontal="center")
        c.execute("SELECT numero_dossier,nom_prenom,telephone,adresse,date_sortie,mois,annee FROM dossiers_sortants ORDER BY id")
        for i,row in enumerate(c.fetchall(),2):
            for j,v in enumerate(row,1): ws2.cell(i,j,v)
        conn.close()
        for ws in [ws1,ws2]:
            for col in ws.columns: ws.column_dimensions[col[0].column_letter].width=18
        wb.save(p); messagebox.showinfo("Succès",f"Excel exporté :\n{p}")

    def _pdf(self):
        if not PDF_OK: messagebox.showerror("Erreur","reportlab non installé"); return
        p=filedialog.asksaveasfilename(defaultextension=".pdf",filetypes=[("PDF","*.pdf")])
        if not p: return
        doc=SimpleDocTemplate(p,pagesize=landscape(A4))
        styles=getSampleStyleSheet()
        el=[]
        t_style=ParagraphStyle("t",fontSize=16,alignment=TA_CENTER,
                               textColor=colors.HexColor("#00b4d8"),fontName="Helvetica-Bold",spaceAfter=8)
        el.append(Paragraph("GJM TECNOLOGIE – Archivage des Dossiers",t_style))
        nom=get_param("nom_entreprise","")
        if nom: el.append(Paragraph(nom,t_style))
        el.append(Paragraph(f"Export du {datetime.date.today().strftime('%d/%m/%Y')}",styles["Normal"]))
        el.append(Spacer(1,0.4*cm))
        conn=get_conn(); c=conn.cursor()
        el.append(Paragraph("DOSSIERS ENTRANTS",ParagraphStyle("h",fontSize=12,
            fontName="Helvetica-Bold",textColor=colors.HexColor("#00b4d8"),spaceAfter=5)))
        c.execute("SELECT numero_dossier,nom_prenom,telephone,commune,nature_dossier,imputer_a,date_entree FROM dossiers_entrants ORDER BY id")
        rows=c.fetchall()
        data=[["N°Dossier","Nom & Prénom","Téléphone","Commune","Nature","Imputer à","Date"]]+[list(r) for r in rows]
        t=Table(data,repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,0),colors.HexColor("#0f3460")),
            ("TEXTCOLOR",(0,0),(-1,0),colors.white),
            ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),
            ("FONTSIZE",(0,0),(-1,-1),8),("ALIGN",(0,0),(-1,-1),"CENTER"),
            ("GRID",(0,0),(-1,-1),0.4,colors.HexColor("#333355")),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.HexColor("#1e2a4a"),colors.HexColor("#16213e")]),
            ("TEXTCOLOR",(0,1),(-1,-1),colors.HexColor("#e0e0e0")),
        ]))
        el.append(t); el.append(Spacer(1,0.5*cm))
        el.append(Paragraph("DOSSIERS SORTANTS",ParagraphStyle("h2",fontSize=12,
            fontName="Helvetica-Bold",textColor=colors.HexColor("#e94560"),spaceAfter=5)))
        c.execute("SELECT numero_dossier,nom_prenom,telephone,adresse,date_sortie FROM dossiers_sortants ORDER BY id")
        rows_s=c.fetchall()
        data_s=[["N°Dossier","Nom & Prénom","Téléphone","Adresse","Date Sortie"]]+[list(r) for r in rows_s]
        ts=Table(data_s,repeatRows=1)
        ts.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,0),colors.HexColor("#e94560")),
            ("TEXTCOLOR",(0,0),(-1,0),colors.white),
            ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),
            ("FONTSIZE",(0,0),(-1,-1),8),("ALIGN",(0,0),(-1,-1),"CENTER"),
            ("GRID",(0,0),(-1,-1),0.4,colors.HexColor("#333355")),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.HexColor("#1e2a4a"),colors.HexColor("#16213e")]),
            ("TEXTCOLOR",(0,1),(-1,-1),colors.HexColor("#e0e0e0")),
        ]))
        el.append(ts); conn.close()
        doc.build(el); messagebox.showinfo("Succès",f"PDF exporté :\n{p}")

    def _word(self):
        if not DOCX_OK: messagebox.showerror("Erreur","python-docx non installé"); return
        p=filedialog.asksaveasfilename(defaultextension=".docx",filetypes=[("Word","*.docx")])
        if not p: return
        doc=DocxDoc()
        h=doc.add_paragraph(); r=h.add_run("GJM TECNOLOGIE – Archivage")
        r.bold=True; r.font.size=Pt(18); r.font.color.rgb=RGBColor(0,0x47,0x60)
        h.alignment=WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Export du {datetime.date.today().strftime('%d/%m/%Y')}")
        conn=get_conn(); c=conn.cursor()
        doc.add_heading("DOSSIERS ENTRANTS",level=1)
        c.execute("SELECT numero_dossier,nom_prenom,telephone,commune,nature_dossier,imputer_a,date_entree FROM dossiers_entrants ORDER BY id")
        rows=c.fetchall()
        if rows:
            ths=["N°Dossier","Nom & Prénom","Téléphone","Commune","Nature","Imputer à","Date"]
            tbl=doc.add_table(rows=1+len(rows),cols=len(ths)); tbl.style="Table Grid"
            for j,h2 in enumerate(ths): tbl.rows[0].cells[j].text=h2
            for i,row in enumerate(rows,1):
                for j,v in enumerate(row): tbl.rows[i].cells[j].text=str(v or "")
        doc.add_heading("DOSSIERS SORTANTS",level=1)
        c.execute("SELECT numero_dossier,nom_prenom,telephone,adresse,date_sortie FROM dossiers_sortants ORDER BY id")
        rows_s=c.fetchall()
        if rows_s:
            ths2=["N°Dossier","Nom & Prénom","Téléphone","Adresse","Date Sortie"]
            tbl2=doc.add_table(rows=1+len(rows_s),cols=len(ths2)); tbl2.style="Table Grid"
            for j,h2 in enumerate(ths2): tbl2.rows[0].cells[j].text=h2
            for i,row in enumerate(rows_s,1):
                for j,v in enumerate(row): tbl2.rows[i].cells[j].text=str(v or "")
        conn.close(); doc.save(p); messagebox.showinfo("Succès",f"Word exporté :\n{p}")

    def _json(self):
        p=filedialog.asksaveasfilename(defaultextension=".json",filetypes=[("JSON","*.json")])
        if not p: return
        conn=get_conn(); c=conn.cursor()
        data={"version":1,"date_backup":str(datetime.datetime.now()),
              "nom_entreprise":get_param("nom_entreprise",""),
              "duree_fin":get_param("duree_fin","")}
        for table,key in [("dossiers_entrants","dossiers_entrants"),
                           ("documents_clients","documents_clients"),
                           ("dossiers_sortants","dossiers_sortants")]:
            c.execute(f"SELECT * FROM {table}")
            cols=[d[0] for d in c.description]
            data[key]=[dict(zip(cols,r)) for r in c.fetchall()]
        conn.close()
        with open(p,"w",encoding="utf-8") as f: json.dump(data,f,ensure_ascii=False,indent=2)
        messagebox.showinfo("Backup réussi",f"Backup JSON enregistré :\n{p}")

    def _restore(self):
        if self.app.role not in ("admin", "superadmin"):
            messagebox.showerror("Acces refuse","Administrateur uniquement."); return
        p = filedialog.askopenfilename(filetypes=[("JSON","*.json")])
        if not p: return
        if not messagebox.askyesno("Confirmation","Remplacer toutes les donnees ?"): return
        try:
            with open(p, encoding="utf-8") as f: data = json.load(f)
        except Exception as e:
            messagebox.showerror("Erreur lecture", str(e)); return
        import time as _t
        for attempt in range(6):
            conn = None
            try:
                conn = sqlite3.connect(str(DB_PATH), timeout=30, check_same_thread=False)
                conn.execute("PRAGMA busy_timeout=30000")
                c = conn.cursor()
                c.execute("DELETE FROM documents_clients")
                c.execute("DELETE FROM dossiers_sortants")
                c.execute("DELETE FROM dossiers_entrants")
                now = str(datetime.datetime.now())
                for row in data.get("dossiers_entrants",[]):
                    row.setdefault("date_creation", now)
                    c.execute("""INSERT OR REPLACE INTO dossiers_entrants
                        (id,numero_dossier,nom_prenom,adresse,telephone,email,
                         ilot_numero,lot_numero,superficie,nom_lotissement,commune,
                         nature_dossier,imputer_a,date_entree,personne_contact,
                         telephone_contact,mois,annee,date_creation)
                        VALUES(:id,:numero_dossier,:nom_prenom,:adresse,:telephone,:email,
                         :ilot_numero,:lot_numero,:superficie,:nom_lotissement,:commune,
                         :nature_dossier,:imputer_a,:date_entree,:personne_contact,
                         :telephone_contact,:mois,:annee,:date_creation)""", row)
                for row in data.get("documents_clients",[]):
                    row.setdefault("date_ajout", now)
                    c.execute("""INSERT OR REPLACE INTO documents_clients
                        (id,dossier_id,nom_fichier,chemin_fichier,type_fichier,date_ajout)
                        VALUES(:id,:dossier_id,:nom_fichier,:chemin_fichier,:type_fichier,:date_ajout)""", row)
                for row in data.get("dossiers_sortants",[]):
                    row.setdefault("date_creation", now)
                    c.execute("""INSERT OR REPLACE INTO dossiers_sortants
                        (id,numero_dossier,nom_prenom,telephone,adresse,date_sortie,mois,annee,date_creation)
                        VALUES(:id,:numero_dossier,:nom_prenom,:telephone,:adresse,
                         :date_sortie,:mois,:annee,:date_creation)""", row)
                conn.commit(); conn.close(); conn = None
                if "nom_entreprise" in data: set_param("nom_entreprise",data["nom_entreprise"])
                if "duree_fin" in data: set_param("duree_fin",data["duree_fin"])
                messagebox.showinfo("Restauration OK","Toutes les donnees ont ete restaurees !")
                return
            except sqlite3.OperationalError as e:
                if conn:
                    try: conn.close()
                    except: pass
                if "locked" in str(e).lower() and attempt < 5:
                    _t.sleep(2); continue
                messagebox.showerror("Base occupee",
                    "Fermez l'application sur les autres PC\npuis reessayez.")
                return
            except Exception as e:
                if conn:
                    try: conn.close()
                    except: pass
                messagebox.showerror("Erreur", str(e)); return

# ══════════════════════════════════════════════════════════════════════════════
# FENETRE OPTIONS
# ══════════════════════════════════════════════════════════════════════════════
class FenetreOptions(tk.Toplevel):
    def __init__(self, parent, callback, niveau="admin"):
        super().__init__(parent)
        self.callback = callback
        self.niveau   = niveau  # "admin" ou "superadmin"
        self.title("Options - Administration")
        self.configure(bg=C_BG)
        self.resizable(True, True)
        self.grab_set()
        w, h = 580, 680
        x = (self.winfo_screenwidth()  - w) // 2
        y = (self.winfo_screenheight() - h) // 2
        self.geometry(f"{w}x{h}+{x}+{y}")
        self.minsize(520, 600)
        self._build()

    def _build(self):
        # En-tête avec indicateur de niveau
        hdr = tk.Frame(self, bg=C_PANEL, pady=15)
        hdr.pack(fill="x")
        if self.niveau == "superadmin":
            tk.Label(hdr, text="Options – GJM Tecnologie (Support)",
                     font=("Segoe UI Black", 13), bg=C_PANEL, fg=C_GOLD).pack()
            tk.Label(hdr, text="Acces complet – Toutes les sections visibles",
                     font=("Segoe UI", 10), bg=C_PANEL, fg=C_SUCCESS).pack(pady=(3,0))
        else:
            tk.Label(hdr, text="Options Administrateur",
                     font=("Segoe UI Black", 14), bg=C_PANEL, fg=C_GOLD).pack()
            tk.Label(hdr, text="Gestion de votre application",
                     font=("Segoe UI", 10), bg=C_PANEL, fg="#aaaaaa").pack(pady=(3,0))

        # Zone scrollable
        canvas = tk.Canvas(self, bg=C_BG, highlightthickness=0)
        sb = ttk.Scrollbar(self, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=sb.set)
        sb.pack(side="right", fill="y")
        canvas.pack(fill="both", expand=True)

        body = tk.Frame(canvas, bg=C_BG, padx=40, pady=20)
        win  = canvas.create_window((0, 0), window=body, anchor="nw")

        def on_resize(e): canvas.itemconfig(win, width=e.width)
        canvas.bind("<Configure>", on_resize)
        body.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.bind_all("<MouseWheel>", lambda e: canvas.yview_scroll(-1*(e.delta//120), "units"))

        # ── SECTION 1 : Parametres generaux ──────────────────────────────────
        tk.Label(body, text="Parametres Generaux",
                 font=("Segoe UI Black", 12), bg=C_BG, fg=C_GOLD).pack(fill="x", pady=(0,12))

        tk.Label(body, text="Nom de l'entreprise :",
                 font=("Segoe UI", 11, "bold"), bg=C_BG, fg=C_GREEN, anchor="w").pack(fill="x", pady=(0,3))
        self.var_nom = tk.StringVar(value=get_param("nom_entreprise", ""))
        make_entry(body, self.var_nom, width=40).pack(fill="x", ipady=6)

        tk.Label(body, text="Date de fin d'utilisation (JJ/MM/AAAA) :",
                 font=("Segoe UI", 11, "bold"), bg=C_BG, fg=C_GREEN, anchor="w").pack(fill="x", pady=(15,3))
        fin = get_param("duree_fin", "")
        if fin:
            try: fin = datetime.datetime.strptime(fin, "%Y-%m-%d").strftime("%d/%m/%Y")
            except: pass
        self.var_fin = tk.StringVar(value=fin)
        make_entry(body, self.var_fin, width=20).pack(anchor="w", ipady=6)

        # Badge statut licence en temps reel
        txt_s, col_s = statut_licence()
        self.badge_lic = tk.Label(body, text=f"  {txt_s}  ",
                                   font=("Segoe UI", 12, "bold"),
                                   bg=col_s, fg="#1a1a2e", padx=8, pady=6)
        self.badge_lic.pack(fill="x", pady=(4,2))
        try:
            fin_d = get_param("duree_fin","")
            if fin_d:
                jours = (datetime.datetime.strptime(fin_d,"%Y-%m-%d").date() - datetime.date.today()).days
                if jours > 0:
                    pct = min(jours/365, 1.0)
                    bg_b = tk.Frame(body, bg=C_PANEL, height=8)
                    bg_b.pack(fill="x", pady=(0,2))
                    bg_b.pack_propagate(False)
                    tk.Frame(bg_b, bg=col_s, height=8).place(relwidth=pct, relheight=1)
                    tk.Label(body, text=f"{jours} jour(s) restant(s)",
                             font=("Segoe UI",10,"italic"), bg=C_BG,
                             fg=col_s, anchor="w").pack(fill="x")
        except: pass


        self.lbl_err = tk.Label(body, text="", bg=C_BG, fg=C_ACCENT, font=("Segoe UI", 11))
        self.lbl_err.pack(pady=3)

        def save():
            nom = self.var_nom.get().strip()
            try:
                d = datetime.datetime.strptime(self.var_fin.get().strip(), "%d/%m/%Y")
            except:
                self.lbl_err.config(text="  Format de date invalide (JJ/MM/AAAA)")
                return
            set_param("nom_entreprise", nom)
            set_param("duree_fin", str(d.date()))
            if hasattr(self,"badge_lic"):
                txt2,col2=statut_licence()
                self.badge_lic.config(text=f"  {txt2}  ",bg=col2)
            self.callback(nom)
            messagebox.showinfo("Sauvegarde", "Options enregistrees !", parent=self)

        make_btn(body, "  Enregistrer les parametres", save,
                 color=C_SUCCESS, fg=C_BG).pack(pady=8, ipady=6, fill="x")

        if self.niveau == "superadmin":
            make_btn(body, "Renouveler la licence",
                     lambda: FenetreRenouvellement(self, callback=self._refresh_badge),
                     color=C_GOLD, fg="#1a1a2e").pack(pady=(0,10), ipady=6, fill="x")

        # ── SECTION 2 : Configuration Reseau ─────────────────────────────────
        tk.Frame(body, bg=C_PANEL, height=2).pack(fill="x", pady=(15, 12))

        tk.Label(body, text="Configuration Reseau",
                 font=("Segoe UI Black", 12), bg=C_BG, fg=C_GOLD).pack(fill="x", pady=(0,8))

        # Statut actuel
        cfg_path   = APP_DIR / "gjm_config.json"
        cfg_actuel = ""
        try:
            import json as _json
            cfg = _json.loads(cfg_path.read_text(encoding="utf-8"))
            cfg_actuel = str(cfg.get("db_path", "")).replace("\\gjm_data.db","").replace("/gjm_data.db","")
        except:
            pass

        if cfg_actuel:
            statut_txt = "Mode reseau actif : " + cfg_actuel
            statut_col = C_SUCCESS
        else:
            statut_txt = "Mode local (pas de reseau configure)"
            statut_col = C_WARNING

        self.lbl_statut = tk.Label(body, text=statut_txt,
                                    font=("Segoe UI", 11, "bold"),
                                    bg=C_BG, fg=statut_col, anchor="w",
                                    wraplength=460, justify="left")
        self.lbl_statut.pack(fill="x", pady=(0,10))

        tk.Label(body, text="Chemin du dossier partage reseau :",
                 font=("Segoe UI", 11, "bold"), bg=C_BG, fg=C_GREEN, anchor="w").pack(fill="x", pady=(0,3))
        tk.Label(body, text="Exemple : \\\\NOM-SERVEUR\\GJM_Partage",
                 font=("Segoe UI", 10), bg=C_BG, fg="#aaaaaa", anchor="w").pack(fill="x")

        var_reseau = tk.StringVar(value=cfg_actuel)
        make_entry(body, var_reseau, width=45).pack(fill="x", ipady=6, pady=(5,3))
        tk.Label(body, text="Laissez vide pour revenir en mode local",
                 font=("Segoe UI", 10, "italic"), bg=C_BG, fg="#888888", anchor="w").pack(fill="x")

        def save_reseau():
            chemin = var_reseau.get().strip()
            cfg_p = APP_DIR / "gjm_config.json"
            if chemin:
                p = Path(chemin)
                if not p.exists():
                    try: p.mkdir(parents=True, exist_ok=True)
                    except: pass
                db_r = p / "gjm_data.db"
                import json as _j
                cfg_p.write_text(_j.dumps({"db_path": str(db_r)}, ensure_ascii=False), encoding="utf-8")
                self.lbl_statut.config(text="Mode reseau actif : " + chemin, fg=C_SUCCESS)
                messagebox.showinfo("Reseau configure",
                    "Base de donnees configuree sur :\n" + str(db_r) +
                    "\n\nRedemarrez l'application pour appliquer.", parent=self)
            else:
                if cfg_p.exists(): cfg_p.unlink()
                self.lbl_statut.config(text="Mode local (pas de reseau configure)", fg=C_WARNING)
                messagebox.showinfo("Mode local",
                    "Configuration reseau supprimee.\n"
                    "Redemarrez l'application pour appliquer.", parent=self)

        make_btn(body, "  Appliquer la configuration reseau",
                 save_reseau, color=C_BLUE).pack(pady=10, ipady=6, fill="x")

        make_btn(body, "  Fermer", self.destroy,
                 color="#333355", fg="#aaaaaa").pack(pady=(5,20), ipady=5, fill="x")


# ══════════════════════════════════════════════════════════════════════════════
# FICHE DOSSIER ENTRANT
# ══════════════════════════════════════════════════════════════════════════════
    def _afficher_liste_users(self):
        for w in self.frame_users.winfo_children(): w.destroy()
        hdr = tk.Frame(self.frame_users, bg=C_PANEL)
        hdr.pack(fill="x", pady=(0,2))
        for txt, w in [("Nom", 22), ("Role", 14), ("Actions", 20)]:
            tk.Label(hdr, text=txt, font=("Segoe UI",10,"bold"),
                     bg=C_PANEL, fg=C_GOLD, width=w, anchor="w").pack(side="left", padx=6, pady=4)
        for nom, info in list(USERS.items()):
            row = tk.Frame(self.frame_users, bg=C_CARD,
                           highlightthickness=1, highlightbackground=C_PANEL)
            row.pack(fill="x", pady=1)
            role_col = C_GOLD if info["role"]=="admin" else C_WHITE
            role_txt = "Admin" if info["role"]=="admin" else "Utilisateur"
            tk.Label(row, text=nom, font=("Segoe UI",11,"bold"),
                     bg=C_CARD, fg=C_WHITE, width=22, anchor="w").pack(side="left", padx=8, pady=5)
            tk.Label(row, text=role_txt, font=("Segoe UI",11),
                     bg=C_CARD, fg=role_col, width=14, anchor="w").pack(side="left")
            make_btn(row, "✏ Modifier", lambda n=nom: self._modifier_user(n),
                     color=C_BLUE).pack(side="left", padx=4, pady=4)
            is_last_admin = (info["role"]=="admin" and
                             sum(1 for u in USERS.values() if u["role"]=="admin") <= 1)
            if not is_last_admin:
                make_btn(row, "🗑 Suppr.", lambda n=nom: self._supprimer_user(n),
                         color=C_ACCENT).pack(side="left", padx=4, pady=4)

    def _ajouter_user(self):
        nom = self.var_new_nom.get().strip()
        pwd = self.var_new_pwd.get().strip()
        role= self.var_new_role.get()
        if not nom:
            self.lbl_user_err.config(text="⚠  Nom obligatoire."); return
        if not pwd:
            self.lbl_user_err.config(text="⚠  Mot de passe obligatoire."); return
        if nom in USERS:
            self.lbl_user_err.config(text="⚠  Ce nom existe deja."); return
        if len(pwd) < 4:
            self.lbl_user_err.config(text="⚠  Mot de passe trop court (4 car. min)."); return
        USERS[nom] = {"password": hash_pwd(pwd), "role": role}
        self.var_new_nom.set(""); self.var_new_pwd.set(""); self.var_new_role.set("user")
        self.lbl_user_err.config(text="✅  " + nom + " ajoute !", fg=C_SUCCESS)
        self._afficher_liste_users()

    def _modifier_user(self, nom):
        fen = tk.Toplevel(self); fen.title("Modifier – " + nom)
        fen.configure(bg=C_BG); fen.grab_set()
        w2,h2=460,340; x2=(fen.winfo_screenwidth()-w2)//2; y2=(fen.winfo_screenheight()-h2)//2
        fen.geometry(f"{w2}x{h2}+{x2}+{y2}")
        tk.Frame(fen, bg=C_PANEL, pady=12).pack(fill="x")
        tk.Label(fen.winfo_children()[0], text="✏  Modifier : "+nom,
                 font=("Segoe UI Black",13), bg=C_PANEL, fg=C_GOLD).pack()
        body2 = tk.Frame(fen, bg=C_BG, padx=30, pady=15); body2.pack(fill="both", expand=True)
        tk.Label(body2, text="Nouveau nom :", font=("Segoe UI",11,"bold"),
                 bg=C_BG, fg=C_GREEN, anchor="w").pack(fill="x", pady=(0,3))
        var_nom2 = tk.StringVar(value=nom)
        make_entry(body2, var_nom2, width=32).pack(fill="x", ipady=5)
        tk.Label(body2, text="Nouveau mot de passe (vide = inchange) :",
                 font=("Segoe UI",11,"bold"), bg=C_BG, fg=C_GREEN, anchor="w").pack(fill="x", pady=(10,3))
        var_pwd2 = tk.StringVar()
        make_entry(body2, var_pwd2, width=32, show="●").pack(fill="x", ipady=5)
        tk.Label(body2, text="Role :", font=("Segoe UI",11,"bold"),
                 bg=C_BG, fg=C_GREEN, anchor="w").pack(fill="x", pady=(10,3))
        var_role2 = tk.StringVar(value=USERS[nom]["role"])
        fr2 = tk.Frame(body2, bg=C_BG); fr2.pack(anchor="w")
        tk.Radiobutton(fr2, text="Utilisateur", variable=var_role2, value="user",
                       bg=C_BG, fg=C_WHITE, selectcolor=C_PANEL, font=("Segoe UI",11)).pack(side="left",padx=(0,12))
        tk.Radiobutton(fr2, text="Administrateur", variable=var_role2, value="admin",
                       bg=C_BG, fg=C_GOLD, selectcolor=C_PANEL, font=("Segoe UI",11)).pack(side="left")
        lbl_e = tk.Label(body2, text="", bg=C_BG, fg=C_ACCENT, font=("Segoe UI",10))
        lbl_e.pack(pady=5)
        def valider():
            nn = var_nom2.get().strip(); np = var_pwd2.get().strip(); nr = var_role2.get()
            if not nn: lbl_e.config(text="⚠  Nom obligatoire."); return
            if nn != nom and nn in USERS: lbl_e.config(text="⚠  Ce nom existe deja."); return
            if np and len(np)<4: lbl_e.config(text="⚠  Mot de passe trop court."); return
            old = USERS.pop(nom)
            USERS[nn] = {"password": hash_pwd(np) if np else old["password"], "role": nr}
            self._afficher_liste_users()
            messagebox.showinfo("Modifie", nn + " mis a jour !", parent=fen)
            fen.destroy()
        make_btn(body2, "💾  Enregistrer", valider, color=C_SUCCESS, fg="#1a1a2e").pack(fill="x",pady=8,ipady=5)
        make_btn(body2, "❌  Annuler", fen.destroy, color="#333355", fg="#aaaaaa").pack(fill="x",ipady=4)

    def _supprimer_user(self, nom):
        msg = "Supprimer l'utilisateur '" + nom + "' ?"
        if not messagebox.askyesno("Confirmer", msg, parent=self): return
        if nom in USERS:
            del USERS[nom]
            self._afficher_liste_users()
            self.lbl_user_err.config(text=nom + " supprime.", fg=C_SUCCESS)


class FicheEntrant(tk.Toplevel):
    def __init__(self, parent, dossier_id=None, callback=None, role="user"):
        super().__init__(parent)
        self.dossier_id=dossier_id; self.callback=callback; self.role=role
        self.data={}
        if dossier_id:
            conn=get_conn(); c=conn.cursor()
            c.execute("SELECT * FROM dossiers_entrants WHERE id=?",(dossier_id,))
            row=c.fetchone(); conn.close()
            if row:
                cols=["id","numero_dossier","nom_prenom","adresse","telephone","email",
                      "ilot_numero","lot_numero","superficie","nom_lotissement","commune",
                      "nature_dossier","imputer_a","date_entree","personne_contact",
                      "telephone_contact","mois","annee","date_creation"]
                self.data=dict(zip(cols,row))
        num=self.data.get("numero_dossier",prochain_numero())
        self.title(f"Dossier {num}")
        self.configure(bg=C_BG)
        self.grab_set()
        w,h=900,720; x=(self.winfo_screenwidth()-w)//2; y=(self.winfo_screenheight()-h)//2
        self.geometry(f"{w}x{h}+{x}+{y}")
        self._build()

    def _build(self):
        # Header
        hdr=tk.Frame(self,bg=C_PANEL,pady=12); hdr.pack(fill="x")
        num=self.data.get("numero_dossier",prochain_numero())
        tk.Label(hdr,text=f"📋  Dossier N°  {num}",font=("Segoe UI Black",15),
                 bg=C_PANEL,fg=C_GOLD).pack(side="left",padx=20)

        # Scroll
        canvas=tk.Canvas(self,bg=C_BG,highlightthickness=0)
        sb=ttk.Scrollbar(self,orient="vertical",command=canvas.yview)
        canvas.configure(yscrollcommand=sb.set)
        sb.pack(side="right",fill="y"); canvas.pack(fill="both",expand=True)
        inner=tk.Frame(canvas,bg=C_BG)
        win=canvas.create_window((0,0),window=inner,anchor="nw")
        canvas.bind("<Configure>",lambda e: canvas.itemconfig(win,width=e.width))
        inner.bind("<Configure>",lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.bind_all("<MouseWheel>",lambda e: canvas.yview_scroll(-1*(e.delta//120),"units"))

        # Champs en grille
        card=make_card(inner,padx=25,pady=20); card.pack(fill="x",padx=20,pady=10)
        tk.Label(card,text="Informations du client",font=("Segoe UI Black",12),
                 bg=C_CARD,fg=C_GREEN).grid(row=0,column=0,columnspan=4,sticky="w",pady=(0,12))

        champs=[
            ("Nom et Prénom *","nom_prenom"),("Adresse","adresse"),
            ("Téléphone","telephone"),("E-mail","email"),
            ("Ilot N°","ilot_numero"),("Lot N°","lot_numero"),
            ("Superficie (m²)","superficie"),("Lotissement","nom_lotissement"),
            ("Commune","commune"),("Nature du dossier","nature_dossier"),
            ("Imputer à","imputer_a"),("Date (JJ/MM/AAAA) *","date_entree"),
            ("Personne à contacter","personne_contact"),("Tél. Contact","telephone_contact"),
        ]
        self.vars={}
        for i,(label,key) in enumerate(champs):
            r,c2=divmod(i,2)
            tk.Label(card,text=label+":",font=("Segoe UI",12,"bold"),
                     bg=C_CARD,fg=C_WHITE,anchor="e",width=22).grid(
                     row=r+1,column=c2*2,padx=8,pady=6,sticky="e")
            var=tk.StringVar(value=self.data.get(key,"") or "")
            self.vars[key]=var
            make_entry(card,var,width=26).grid(row=r+1,column=c2*2+1,padx=8,pady=6,ipady=4)

        # Section documents
        card2=make_card(inner,padx=25,pady=15); card2.pack(fill="x",padx=20,pady=5)
        tk.Label(card2,text="📎  Documents du client",font=("Segoe UI Black",12),
                 bg=C_CARD,fg=C_GREEN).pack(anchor="w",pady=(0,10))

        barre=tk.Frame(card2,bg=C_CARD); barre.pack(fill="x")
        make_btn(barre,"📁  Insérer document(s)",self._inserer,color=C_BLUE).pack(side="left",padx=(0,10))
        make_btn(barre,"📱  Résumé",self._resume,color=C_SUCCESS,fg=C_BG).pack(side="left",padx=5)
        if self.role=="admin" and self.dossier_id:
            make_btn(barre,"🗑  Supprimer",self._supprimer,color=C_ACCENT).pack(side="right")

        cols_d=("Nom du fichier","Type","Date d'ajout")
        frame_d=tk.Frame(card2,bg=C_CARD,highlightthickness=1,highlightbackground=C_PANEL)
        frame_d.pack(fill="x",pady=8)
        self.tree_docs=ttk.Treeview(frame_d,columns=cols_d,show="headings",
                                     style="Modern.Treeview",height=5)
        for col in cols_d:
            self.tree_docs.heading(col,text=col)
            self.tree_docs.column(col,width=250)
        sb_d=ttk.Scrollbar(frame_d,orient="vertical",command=self.tree_docs.yview)
        self.tree_docs.configure(yscrollcommand=sb_d.set)
        sb_d.pack(side="right",fill="y"); self.tree_docs.pack(fill="x")
        self.tree_docs.bind("<Double-1>",self._ouvrir_doc)
        self._charger_docs()

        # Boutons bas
        bas=tk.Frame(inner,bg=C_BG,pady=15); bas.pack()
        make_btn(bas,"💾  Enregistrer",self._save,color=C_SUCCESS,fg=C_BG).pack(side="left",padx=10,ipady=5)
        make_btn(bas,"❌  Fermer",self.destroy,color="#333355",fg="#aaaaaa").pack(side="left",padx=10,ipady=5)

    def _charger_docs(self):
        self.tree_docs.delete(*self.tree_docs.get_children())
        if not self.dossier_id: return
        conn=get_conn(); c=conn.cursor()
        c.execute("SELECT id,nom_fichier,type_fichier,date_ajout FROM documents_clients WHERE dossier_id=?",(self.dossier_id,))
        for i,row in enumerate(c.fetchall()):
            self.tree_docs.insert("","end",iid=str(row[0]),values=(row[1],row[2],row[3]),
                tags=("pair" if i%2==0 else "impair",))
        conn.close()

    def _inserer(self):
        if not self.dossier_id:
            did=self._save(ret=True)
            if not did: return
        fichiers=filedialog.askopenfilenames(title="Sélectionner les documents",
            filetypes=[("Tous","*.*"),("PDF","*.pdf"),("Images","*.png *.jpg *.jpeg"),
                       ("Word","*.docx *.doc"),("DWG","*.dwg *.dxf"),("Excel","*.xlsx")])
        if not fichiers: return
        nom=self.vars["nom_prenom"].get().strip().replace(" ","_")
        conn=get_conn(); c=conn.cursor()
        c.execute("SELECT numero_dossier FROM dossiers_entrants WHERE id=?",(self.dossier_id,))
        r=c.fetchone(); conn.close()
        num=r[0] if r else str(self.dossier_id)
        dest_dir=DOCS_DIR/f"{nom}_{num}"; dest_dir.mkdir(parents=True,exist_ok=True)
        conn=get_conn(); c=conn.cursor()
        for f_src in fichiers:
            nom_f=Path(f_src).name
            dest=dest_dir/nom_f; shutil.copy2(f_src,dest)
            ext=Path(f_src).suffix.upper().lstrip(".")
            c.execute("INSERT INTO documents_clients(dossier_id,nom_fichier,chemin_fichier,type_fichier) VALUES(?,?,?,?)",
                      (self.dossier_id,nom_f,str(dest),ext))
        conn.commit(); conn.close()
        self._charger_docs()
        messagebox.showinfo("Documents ajoutés",f"{len(fichiers)} document(s) ajouté(s) !")

    def _ouvrir_doc(self,e):
        sel=self.tree_docs.focus()
        if not sel: return
        conn=get_conn(); c=conn.cursor()
        c.execute("SELECT chemin_fichier FROM documents_clients WHERE id=?",(int(sel),))
        r=c.fetchone(); conn.close()
        if r and Path(r[0]).exists(): ouvrir_fichier(r[0])
        else: messagebox.showerror("Erreur","Fichier introuvable.")

    def _resume(self):
        num  = self.data.get("numero_dossier","N/A")
        tel  = self.vars['telephone'].get().strip()
        docs = []
        if self.dossier_id:
            conn=get_conn(); c=conn.cursor()
            c.execute("SELECT nom_fichier,type_fichier FROM documents_clients WHERE dossier_id=?",(self.dossier_id,))
            docs=c.fetchall(); conn.close()

        nom_ent = get_param('nom_entreprise','GJM Tecnologie')

        txt = (
            f"📋 *RÉSUMÉ DOSSIER – {nom_ent.upper()}*\n"
            f"{'─'*40}\n\n"
            f"🔢 *N° Dossier   :* {num}\n"
            f"👤 *Nom & Prénom :* {self.vars['nom_prenom'].get()}\n"
            f"📞 *Téléphone    :* {tel}\n"
            f"📧 *E-mail       :* {self.vars['email'].get()}\n"
            f"🏠 *Adresse      :* {self.vars['adresse'].get()}\n\n"
            f"🏘 *Ilot N°      :* {self.vars['ilot_numero'].get()}\n"
            f"🏘 *Lot N°       :* {self.vars['lot_numero'].get()}\n"
            f"📐 *Superficie   :* {self.vars['superficie'].get()} m²\n"
            f"🏙 *Lotissement  :* {self.vars['nom_lotissement'].get()}\n"
            f"🌍 *Commune      :* {self.vars['commune'].get()}\n\n"
            f"📁 *Nature dossier :* {self.vars['nature_dossier'].get()}\n"
            f"👔 *Imputer à      :* {self.vars['imputer_a'].get()}\n"
            f"📅 *Date entrée    :* {self.vars['date_entree'].get()}\n\n"
            f"👤 *Personne à contacter :* {self.vars['personne_contact'].get()}\n"
            f"📞 *Tél. Contact         :* {self.vars['telephone_contact'].get()}\n\n"
        )
        if docs:
            txt += f"📎 *Documents fournis ({len(docs)}) :*\n"
            for d in docs: txt += f"  • {d[0]}  ({d[1]})\n"
        else:
            txt += "📎 *Documents :* Aucun document enregistré.\n"
        txt += f"\n{'─'*40}\n_{nom_ent}_"

        # Nettoyer le numéro pour WhatsApp (chiffres seulement)
        tel_clean = "".join(ch for ch in tel if ch.isdigit())
        # Si le numéro commence par 0, ajouter indicatif Côte d'Ivoire 225
        if tel_clean.startswith("0") and len(tel_clean) <= 10:
            tel_clean = "225" + tel_clean[1:]
        elif not tel_clean.startswith("225") and len(tel_clean) <= 10:
            tel_clean = "225" + tel_clean

        import urllib.parse
        msg_encode = urllib.parse.quote(txt)
        url_whatsapp = f"https://wa.me/{tel_clean}?text={msg_encode}"

        # Fenêtre résumé
        fen = tk.Toplevel(self)
        fen.title("Résumé – " + num)
        fen.configure(bg=C_BG)
        fen.geometry("620x580")
        x=(fen.winfo_screenwidth()-620)//2; y=(fen.winfo_screenheight()-580)//2
        fen.geometry(f"620x580+{x}+{y}")
        fen.grab_set()

        # Header
        hdr=tk.Frame(fen,bg=C_PANEL,pady=12); hdr.pack(fill="x")
        tk.Label(hdr,text=f"📋  Résumé Dossier {num}",font=("Segoe UI Black",14),
                 bg=C_PANEL,fg=C_GOLD).pack()

        # Texte
        frame_t=tk.Frame(fen,bg=C_BG); frame_t.pack(fill="both",expand=True,padx=12,pady=8)
        t=tk.Text(frame_t,wrap="word",font=("Segoe UI",13),bg=C_CARD,fg=C_WHITE,
                  insertbackground=C_WHITE,padx=15,pady=12,relief="flat",
                  spacing1=3,spacing2=2)
        sb_t=ttk.Scrollbar(frame_t,orient="vertical",command=t.yview)
        t.configure(yscrollcommand=sb_t.set)
        sb_t.pack(side="right",fill="y"); t.pack(fill="both",expand=True)
        t.insert("1.0",txt)
        t.config(state="disabled")

        # Boutons
        bas=tk.Frame(fen,bg=C_BG,pady=10); bas.pack()

        def envoyer_whatsapp():
            import webbrowser
            webbrowser.open(url_whatsapp)

        make_btn(bas,"📲  Envoyer via WhatsApp",envoyer_whatsapp,
                 color="#25D366",fg=C_WHITE).pack(side="left",padx=10,ipady=6)
        make_btn(bas,"📋  Copier le texte",
                 lambda: (fen.clipboard_clear(), fen.clipboard_append(txt),
                          messagebox.showinfo("Copié","Texte copié !",parent=fen)),
                 color=C_BLUE).pack(side="left",padx=10,ipady=6)
        make_btn(bas,"❌  Fermer",fen.destroy,color="#333355",fg="#aaaaaa").pack(side="left",padx=10,ipady=6)

    def _save(self, ret=False):
        # Bloquer si aucune durée définie
        if not get_param("duree_fin",""):
            messagebox.showerror("Application non configurée",
                "L'administrateur doit définir la durée d'utilisation\n"
                "dans OPTIONS avant de pouvoir enregistrer des données.",
                parent=self)
            return None
        # Bloquer si durée expirée
        fin = get_param("duree_fin","")
        if fin and datetime.date.today() > datetime.datetime.strptime(fin,"%Y-%m-%d").date():
            messagebox.showerror("Licence expirée",
                "La durée d'utilisation est expirée.\n"
                "Contactez l'administrateur pour renouveler.",
                parent=self)
            return None
        nom=self.vars["nom_prenom"].get().strip()
        date_s=self.vars["date_entree"].get().strip()
        if not nom:
            messagebox.showerror("Champ requis","Nom et prénom obligatoire.",parent=self); return None
        try: d=datetime.datetime.strptime(date_s,"%d/%m/%Y"); mois,annee=d.month,d.year
        except: mois=annee=None
        conn=get_conn(); c=conn.cursor()
        v={k:self.vars[k].get().strip() for k in self.vars}
        if not self.dossier_id:
            num=prochain_numero()
            c.execute("""INSERT INTO dossiers_entrants
                (numero_dossier,nom_prenom,adresse,telephone,email,ilot_numero,lot_numero,
                 superficie,nom_lotissement,commune,nature_dossier,imputer_a,date_entree,
                 personne_contact,telephone_contact,mois,annee) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (num,v["nom_prenom"],v["adresse"],v["telephone"],v["email"],v["ilot_numero"],
                 v["lot_numero"],v["superficie"],v["nom_lotissement"],v["commune"],
                 v["nature_dossier"],v["imputer_a"],v["date_entree"],v["personne_contact"],
                 v["telephone_contact"],mois,annee))
            new_id=c.lastrowid; conn.commit(); conn.close()
            if ret:
                self.dossier_id=new_id
                conn2=get_conn(); c2=conn2.cursor()
                c2.execute("SELECT numero_dossier FROM dossiers_entrants WHERE id=?",(new_id,))
                r=c2.fetchone(); conn2.close()
                if r: self.data["numero_dossier"]=r[0]
                return new_id
            messagebox.showinfo("Enregistré","Dossier enregistré !",parent=self)
            if self.callback: self.callback()
            self.destroy()
        else:
            c.execute("""UPDATE dossiers_entrants SET nom_prenom=?,adresse=?,telephone=?,email=?,
                ilot_numero=?,lot_numero=?,superficie=?,nom_lotissement=?,commune=?,nature_dossier=?,
                imputer_a=?,date_entree=?,personne_contact=?,telephone_contact=?,mois=?,annee=?
                WHERE id=?""",
                (v["nom_prenom"],v["adresse"],v["telephone"],v["email"],v["ilot_numero"],
                 v["lot_numero"],v["superficie"],v["nom_lotissement"],v["commune"],
                 v["nature_dossier"],v["imputer_a"],v["date_entree"],v["personne_contact"],
                 v["telephone_contact"],mois,annee,self.dossier_id))
            conn.commit(); conn.close()
            if not ret:
                messagebox.showinfo("Modifié","Dossier mis à jour !",parent=self)
                if self.callback: self.callback()

    def _supprimer(self):
        if self.role not in ("admin","superadmin"): messagebox.showerror("Accès refusé","Administrateur uniquement."); return
        if messagebox.askyesno("Confirmer","Supprimer ce dossier et tous ses documents ?",parent=self):
            conn=get_conn(); c=conn.cursor()
            c.execute("DELETE FROM documents_clients WHERE dossier_id=?",(self.dossier_id,))
            c.execute("DELETE FROM dossiers_entrants WHERE id=?",(self.dossier_id,))
            conn.commit(); conn.close()
            if self.callback: self.callback()
            self.destroy()

# ══════════════════════════════════════════════════════════════════════════════
# FICHE DOSSIER SORTANT
# ══════════════════════════════════════════════════════════════════════════════
class FicheSortant(tk.Toplevel):
    def __init__(self, parent, dossier_id=None, callback=None, role="user"):
        super().__init__(parent)
        self.dossier_id=dossier_id; self.callback=callback; self.role=role; self.data={}
        if dossier_id:
            conn=get_conn(); c=conn.cursor()
            c.execute("SELECT * FROM dossiers_sortants WHERE id=?",(dossier_id,))
            row=c.fetchone(); conn.close()
            if row:
                cols=["id","numero_dossier","nom_prenom","telephone","adresse","date_sortie","mois","annee","date_creation"]
                self.data=dict(zip(cols,row))
        self.title("Dossier Sortant"); self.configure(bg=C_BG)
        self.grab_set()
        w,h=620,460; x=(self.winfo_screenwidth()-w)//2; y=(self.winfo_screenheight()-h)//2
        self.geometry(f"{w}x{h}+{x}+{y}")
        self._build()

    def _build(self):
        hdr=tk.Frame(self,bg=C_ACCENT,pady=12); hdr.pack(fill="x")
        tk.Label(hdr,text="📤  Dossier Sortant",font=("Segoe UI Black",14),
                 bg=C_ACCENT,fg=C_WHITE).pack()

        body=tk.Frame(self,bg=C_BG,padx=35,pady=20); body.pack(fill="both",expand=True)

        tk.Label(body,text="N° Dossier (depuis entrants) :",font=("Segoe UI",11,"bold"),
                 bg=C_BG,fg=C_GREEN,anchor="w").pack(fill="x",pady=(0,3))
        conn=get_conn(); c=conn.cursor()
        c.execute("SELECT numero_dossier,nom_prenom FROM dossiers_entrants ORDER BY id")
        rows=c.fetchall(); conn.close()
        nums=[f"{r[0]}  –  {r[1]}" for r in rows]
        self.var_num=tk.StringVar()
        cb=ttk.Combobox(body,textvariable=self.var_num,values=nums,state="readonly",font=("Segoe UI",11))
        cb.pack(fill="x",ipady=6,pady=(0,12))
        if self.data.get("numero_dossier"):
            for n in nums:
                if n.startswith(self.data["numero_dossier"]): cb.set(n); break
        cb.bind("<<ComboboxSelected>>",self._auto)

        champs=[("Nom & Prénom","nom_prenom"),("Téléphone","telephone"),
                ("Adresse","adresse"),("Date Sortie (JJ/MM/AAAA) *","date_sortie")]
        self.vars={}
        for label,key in champs:
            tk.Label(body,text=label+":",font=("Segoe UI",11,"bold"),
                     bg=C_BG,fg=C_WHITE,anchor="w").pack(fill="x",pady=(8,2))
            var=tk.StringVar(value=self.data.get(key,"") or "")
            self.vars[key]=var
            make_entry(body,var,width=40).pack(fill="x",ipady=5)

        bas=tk.Frame(self,bg=C_BG,pady=12); bas.pack()
        make_btn(bas,"💾  Enregistrer",self._save,color=C_SUCCESS,fg=C_BG).pack(side="left",padx=10,ipady=5)
        if self.role=="admin" and self.dossier_id:
            make_btn(bas,"🗑  Supprimer",self._del,color=C_ACCENT).pack(side="left",padx=10,ipady=5)
        make_btn(bas,"❌  Fermer",self.destroy,color="#333355",fg="#aaaaaa").pack(side="left",padx=10,ipady=5)

    def _auto(self,e):
        sel=self.var_num.get(); num=sel.split("–")[0].strip()
        conn=get_conn(); c=conn.cursor()
        c.execute("SELECT nom_prenom,telephone,adresse FROM dossiers_entrants WHERE numero_dossier=?",(num,))
        r=c.fetchone(); conn.close()
        if r:
            self.vars["nom_prenom"].set(r[0] or "")
            self.vars["telephone"].set(r[1] or "")
            self.vars["adresse"].set(r[2] or "")

    def _save(self):
        # Bloquer si aucune durée définie ou expirée
        if not get_param("duree_fin",""):
            messagebox.showerror("Non configuré",
                "L'administrateur doit définir la durée d'utilisation dans OPTIONS.",parent=self)
            return
        fin=get_param("duree_fin","")
        if fin and datetime.date.today() > datetime.datetime.strptime(fin,"%Y-%m-%d").date():
            messagebox.showerror("Licence expirée","Contactez l'administrateur.",parent=self)
            return
        sel=self.var_num.get(); num=sel.split("–")[0].strip() if "–" in sel else ""
        if not num: messagebox.showerror("Requis","Sélectionnez un N° de dossier.",parent=self); return
        try: d=datetime.datetime.strptime(self.vars["date_sortie"].get().strip(),"%d/%m/%Y"); mois,annee=d.month,d.year
        except: messagebox.showerror("Date invalide","Format JJ/MM/AAAA",parent=self); return
        conn=get_conn(); c=conn.cursor()
        if not self.dossier_id:
            c.execute("INSERT INTO dossiers_sortants(numero_dossier,nom_prenom,telephone,adresse,date_sortie,mois,annee) VALUES(?,?,?,?,?,?,?)",
                (num,self.vars["nom_prenom"].get(),self.vars["telephone"].get(),
                 self.vars["adresse"].get(),self.vars["date_sortie"].get(),mois,annee))
        else:
            c.execute("UPDATE dossiers_sortants SET numero_dossier=?,nom_prenom=?,telephone=?,adresse=?,date_sortie=?,mois=?,annee=? WHERE id=?",
                (num,self.vars["nom_prenom"].get(),self.vars["telephone"].get(),
                 self.vars["adresse"].get(),self.vars["date_sortie"].get(),mois,annee,self.dossier_id))
        conn.commit(); conn.close()
        messagebox.showinfo("Enregistré","Dossier sortant enregistré !",parent=self)
        if self.callback: self.callback()
        self.destroy()

    def _del(self):
        if messagebox.askyesno("Confirmer","Supprimer ce dossier sortant ?",parent=self):
            conn=get_conn(); c=conn.cursor()
            c.execute("DELETE FROM dossiers_sortants WHERE id=?",(self.dossier_id,))
            conn.commit(); conn.close()
            if self.callback: self.callback()
            self.destroy()

if __name__ == "__main__":
    app = AppGJM()
    app.mainloop()